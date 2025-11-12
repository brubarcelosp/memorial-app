from flask import Flask, render_template, request, send_file, jsonify, session
import os
import io
import re
import math
from datetime import datetime
from werkzeug.utils import secure_filename
import tempfile
import shutil

# Importar funções do módulo de processamento
from memorial_processor import (
    _build_memorial_resumo_doc_web, _build_solicitacao_analise_doc_web,
    build_unif_desm_doc_web, build_condominio_loteamento_doc_web,
    build_excel_fracao_ideal_web, build_excel_vertices_web
)

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-this-in-production'
app.config['UPLOAD_FOLDER'] = tempfile.mkdtemp()
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Criar diretórios necessários
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs('static/uploads', exist_ok=True)
os.makedirs('static/images', exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/upload', methods=['POST'])
def enviar_arquivos():
    """Endpoint para upload de arquivos"""
    if 'files' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400
    
    arquivos = request.files.getlist('files')
    arquivos_enviados = {}
    
    for arquivo in arquivos:
        if arquivo.filename == '':
            continue
        
        if arquivo and arquivo_permitido(arquivo.filename):
            nome_arquivo = secure_filename(arquivo.filename)
            caminho_arquivo = os.path.join(app.config['UPLOAD_FOLDER'], nome_arquivo)
            arquivo.save(caminho_arquivo)
            
            with open(caminho_arquivo, 'rb') as f:
                arquivos_enviados[nome_arquivo] = f.read()
    
    # Armazenar na sessão
    session['uploaded_files'] = arquivos_enviados
    
    return jsonify({
        'success': True,
        'count': len(arquivos_enviados),
        'files': list(arquivos_enviados.keys())
    })

def arquivo_permitido(nome_arquivo):
    return '.' in nome_arquivo and \
           nome_arquivo.rsplit('.', 1)[1].lower() in ['html', 'htm', 'txt']

def arquivo_imagem_permitido(nome_arquivo):
    return '.' in nome_arquivo and \
           nome_arquivo.rsplit('.', 1)[1].lower() in ['png', 'jpg', 'jpeg', 'gif', 'bmp']

@app.route('/api/upload-image', methods=['POST'])
def enviar_imagem():
    """Endpoint para upload de imagens (marca d'água, logos, etc.)"""
    if 'file' not in request.files:
        return jsonify({'error': 'Nenhum arquivo enviado'}), 400
    
    arquivo = request.files['file']
    tipo_imagem = request.form.get('type', 'marca_dagua')  # marca_dagua, logo_cabecalho, logo_rodape
    
    if arquivo.filename == '':
        return jsonify({'error': 'Nome de arquivo vazio'}), 400
    
    if arquivo and arquivo_imagem_permitido(arquivo.filename):
        # Definir nome do arquivo baseado no tipo
        if tipo_imagem == 'marca_dagua':
            nome_arquivo = 'marca_dagua.png'
        elif tipo_imagem == 'logo_cabecalho':
            nome_arquivo = 'logo_cabecalho.png'
        elif tipo_imagem == 'logo_rodape':
            nome_arquivo = 'logo_rodape.png'
        else:
            nome_arquivo = secure_filename(arquivo.filename)
        
        caminho_arquivo = os.path.join('static/images', nome_arquivo)
        
        # Garantir que o diretório existe
        os.makedirs('static/images', exist_ok=True)
        
        # Salvar arquivo
        arquivo.save(caminho_arquivo)
        
        return jsonify({
            'success': True,
            'filename': nome_arquivo,
            'path': caminho_arquivo,
            'message': f'Imagem {tipo_imagem} salva com sucesso!'
        })
    
    return jsonify({'error': 'Tipo de arquivo não permitido. Use PNG, JPG, JPEG, GIF ou BMP'}), 400

@app.route('/api/generate', methods=['POST'])
def gerar_documento():
    """Endpoint principal para gerar documentos"""
    try:
        dados = request.get_json()
        modo = dados.get('tipo_emp')
        
        # Recuperar arquivos da sessão
        arquivos_enviados = session.get('uploaded_files', {})
        
        # Criar objeto de contexto com os valores do formulário
        dados_formulario = ContextoDadosFormulario(dados)
        
        # Criar diretório temporário para output
        diretorio_saida = tempfile.mkdtemp()
        
        try:
            if modo == 'memorial_resumo':
                caminho_saida = gerar_memorial_resumo(dados_formulario, diretorio_saida)
            elif modo == 'solicitacao_analise':
                caminho_saida = gerar_solicitacao_analise(dados_formulario, diretorio_saida)
            elif modo in ('unificacao', 'desmembramento', 'unif_desm'):
                caminho_saida = gerar_unif_desm(dados_formulario, arquivos_enviados, modo, diretorio_saida)
            else:  # condominio ou loteamento
                caminho_saida = gerar_condominio_loteamento(
                    dados_formulario, arquivos_enviados, modo, diretorio_saida
                )
            
            # Verificar se o arquivo foi gerado corretamente
            if not caminho_saida or not os.path.exists(caminho_saida):
                raise Exception(f"Arquivo não foi gerado: {caminho_saida}")
            
            tamanho_arquivo = os.path.getsize(caminho_saida)
            if tamanho_arquivo == 0:
                raise Exception(f"Arquivo gerado está vazio: {caminho_saida}")
            
            # Verificar conteúdo do arquivo antes de mover
            try:
                from docx import Document
                doc = Document(caminho_saida)
                contagem_paragrafos = len(doc.paragraphs)
                print(f"📄 Arquivo gerado: {caminho_saida}, Tamanho: {tamanho_arquivo} bytes, Parágrafos: {contagem_paragrafos}")
                if contagem_paragrafos == 0:
                    print(f"⚠️ AVISO: Arquivo tem 0 parágrafos antes de mover!")
            except Exception as e:
                print(f"⚠️ AVISO: Erro ao verificar arquivo antes de mover: {e}")
            
            # Mover arquivo para o diretório de uploads para download
            nome_arquivo = os.path.basename(caminho_saida)
            caminho_destino = os.path.join(app.config['UPLOAD_FOLDER'], nome_arquivo)
            
            # Garantir que o diretório existe
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            
            shutil.move(caminho_saida, caminho_destino)
            
            # Verificar novamente após mover
            if not os.path.exists(caminho_destino):
                raise Exception(f"Arquivo não foi movido corretamente para: {caminho_destino}")
            
            tamanho_final = os.path.getsize(caminho_destino)
            if tamanho_final == 0:
                raise Exception(f"Arquivo ficou vazio após mover: {caminho_destino}")
            
            # Verificar conteúdo do arquivo após mover
            try:
                from docx import Document
                doc = Document(caminho_destino)
                contagem_paragrafos = len(doc.paragraphs)
                print(f"📄 Arquivo movido: {caminho_destino}, Tamanho: {tamanho_final} bytes, Parágrafos: {contagem_paragrafos}")
                if contagem_paragrafos == 0:
                    print(f"⚠️ AVISO: Arquivo tem 0 parágrafos após mover!")
            except Exception as e:
                print(f"⚠️ AVISO: Erro ao verificar arquivo após mover: {e}")
            
            # Copiar também para Downloads para garantir que o arquivo está acessível
            caminho_downloads = os.path.expanduser(f'~/Downloads/{nome_arquivo}')
            try:
                shutil.copy2(caminho_destino, caminho_downloads)
                print(f"📥 Arquivo também copiado para: {caminho_downloads}")
                # Verificar se a cópia está OK
                from docx import Document
                doc_verificacao = Document(caminho_downloads)
                print(f"📄 Cópia em Downloads - Parágrafos: {len(doc_verificacao.paragraphs)}")
            except Exception as e:
                print(f"⚠️ Aviso: Não foi possível copiar para Downloads: {e}")
            
            return jsonify({
                'success': True,
                'filename': nome_arquivo,
                'download_url': f'/api/download/{nome_arquivo}',
                'file_size': tamanho_final,
                'downloads_path': caminho_downloads if 'caminho_downloads' in locals() else None
            })
        finally:
            # Limpar diretório temporário se necessário
            try:
                if os.path.exists(diretorio_saida):
                    shutil.rmtree(diretorio_saida)
            except:
                pass
    
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500

@app.route('/api/download/<filename>')
def baixar_arquivo(nome_arquivo):
    """Endpoint para download de arquivos gerados"""
    # Buscar o arquivo no diretório de uploads
    caminho_arquivo = os.path.join(app.config['UPLOAD_FOLDER'], nome_arquivo)
    
    if not os.path.exists(caminho_arquivo):
        return jsonify({'error': 'Arquivo não encontrado'}), 404
    
    # Verificar se o arquivo não está vazio
    tamanho_arquivo = os.path.getsize(caminho_arquivo)
    if tamanho_arquivo == 0:
        return jsonify({'error': 'Arquivo está vazio'}), 500
    
    # Verificar se é um arquivo .docx válido
    try:
        from docx import Document
        doc = Document(caminho_arquivo)
        if len(doc.paragraphs) == 0:
            print(f"⚠️ AVISO: Arquivo {nome_arquivo} tem 0 parágrafos!")
    except Exception as e:
        print(f"⚠️ AVISO: Erro ao verificar arquivo {nome_arquivo}: {e}")
    
    return send_file(
        caminho_arquivo, 
        as_attachment=True,
        download_name=nome_arquivo,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/api/generate-excel', methods=['POST'])
def gerar_excel():
    """Endpoint para gerar planilhas Excel"""
    try:
        dados = request.get_json()
        modo = dados.get('tipo_emp')
        arquivos_enviados = session.get('uploaded_files', {})
        dados_formulario = ContextoDadosFormulario(dados)
        
        diretorio_saida = tempfile.mkdtemp()
        
        try:
            if modo == 'condominio':
                # Excel de fração ideal
                caminho_saida = gerar_excel_fracao_ideal(dados_formulario, arquivos_enviados, diretorio_saida)
            elif modo in ('unificacao', 'desmembramento', 'unif_desm'):
                # Excel de vértices
                caminho_saida = gerar_excel_vertices(dados_formulario, arquivos_enviados, modo, diretorio_saida)
            else:
                return jsonify({'error': 'Tipo não suporta Excel'}), 400
            
            # Mover arquivo para o diretório de uploads para download
            nome_arquivo = os.path.basename(caminho_saida)
            caminho_destino = os.path.join(app.config['UPLOAD_FOLDER'], nome_arquivo)
            shutil.move(caminho_saida, caminho_destino)
            
            return jsonify({
                'success': True,
                'filename': nome_arquivo,
                'download_url': f'/api/download/{nome_arquivo}'
            })
        finally:
            # Limpar diretório temporário se necessário
            try:
                if os.path.exists(diretorio_saida):
                    shutil.rmtree(diretorio_saida)
            except:
                pass
    
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500

# Classes auxiliares para simular os widgets
class ContextoDadosFormulario:
    """Simula os widgets do ipywidgets usando dados do formulário"""
    def __init__(self, dados):
        self.dados = dados
    
    def obter(self, chave, padrao=''):
        return self.dados.get(chave, padrao)
    
    @property
    def tipo_emp(self):
        return self.obter('tipo_emp', 'condominio')
    
    @property
    def nome_emp(self):
        return self.obter('nome_emp', '')
    
    @property
    def endereco_emp(self):
        return self.obter('endereco_emp', '')
    
    @property
    def bairro_emp(self):
        return self.obter('bairro_emp', '')
    
    @property
    def cidade_emp(self):
        return self.obter('cidade_emp', '')
    
    @property
    def area_total_emp(self):
        return self.obter('area_total_emp', '')
    
    @property
    def perimetro_emp(self):
        return self.obter('perimetro_emp', '')
    
    @property
    def matricula_emp(self):
        return self.obter('matricula_emp', '')
    
    @property
    def num_lotes_emp(self):
        return int(self.obter('num_lotes_emp', 0) or 0)
    
    @property
    def area_tot_priv_emp(self):
        return self.obter('area_tot_priv_emp', '')
    
    @property
    def area_tot_cond_emp(self):
        return self.obter('area_tot_cond_emp', '')
    
    @property
    def ane_drop(self):
        return self.obter('ane_drop', 'Não')
    
    @property
    def ane_largura(self):
        return self.obter('ane_largura', '')
    
    @property
    def coord_fmt(self):
        class FormatoCoord:
            value = self.obter('coord_fmt', 'utm')
        return FormatoCoord()
    
    @property
    def tipo_proj_resumo(self):
        class TipoProjeto:
            value = self.obter('tipo_proj_resumo', 'condominio')
        return TipoProjeto()
    
    @property
    def usos_multi(self):
        class Usos:
            value = self.obter('usos_multi', [])
        return Usos()
    
    @property
    def topografia(self):
        class Topografia:
            value = self.obter('topografia', 'Acentuada')
        return Topografia()
    
    @property
    def has_ai(self):
        class TemAI:
            value = self.obter('has_ai', False)
        return TemAI()
    
    @property
    def has_restricao(self):
        class TemRestricao:
            value = self.obter('has_restricao', False)
        return TemRestricao()

# Funções de geração de documentos
def gerar_memorial_resumo(dados_formulario, diretorio_saida):
    """Gera memorial resumo"""
    return _build_memorial_resumo_doc_web(dados_formulario, diretorio_saida)

def gerar_solicitacao_analise(dados_formulario, diretorio_saida):
    """Gera solicitação de análise"""
    return _build_solicitacao_analise_doc_web(dados_formulario, diretorio_saida)

def gerar_unif_desm(dados_formulario, arquivos_enviados, modo, diretorio_saida):
    """Gera documentos de unificação/desmembramento"""
    return build_unif_desm_doc_web(dados_formulario, arquivos_enviados, modo, diretorio_saida)

def gerar_condominio_loteamento(dados_formulario, arquivos_enviados, modo, diretorio_saida):
    """Gera documentos de condomínio ou loteamento"""
    return build_condominio_loteamento_doc_web(dados_formulario, arquivos_enviados, modo, diretorio_saida)

def gerar_excel_fracao_ideal(dados_formulario, arquivos_enviados, diretorio_saida):
    """Gera Excel de fração ideal"""
    return build_excel_fracao_ideal_web(dados_formulario, arquivos_enviados, diretorio_saida)

def gerar_excel_vertices(dados_formulario, arquivos_enviados, modo, diretorio_saida):
    """Gera Excel de vértices"""
    return build_excel_vertices_web(dados_formulario, arquivos_enviados, modo, diretorio_saida)

if __name__ == '__main__':
    # Para desenvolvimento local
    port = int(os.environ.get('PORT', 5001))
    app.run(debug=True, host='0.0.0.0', port=port)

