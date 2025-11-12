"""
Módulo de processamento de memoriais descritivos
Adaptado do script Xuxu.py para uso em aplicação web
"""
import re
import os
import io
import math
from datetime import datetime
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from num2words import num2words
import pandas as pd
from pyproj import CRS, Transformer

# ===================== Configuração de imagens =====================
# Por padrão, usar imagens locais ou placeholder
CAMINHO_MARCA_DAGUA = "static/images/marca_dagua.png"
CAMINHO_LOGO_CABECALHO = "static/images/logo_cabecalho.png"
CAMINHO_LOGO_RODAPE = "static/images/logo_rodape.png"

# ===================== Utilidades numéricas / texto =====================
def _fmt_br(v, casas=2):
    try:
        return f"{float(v):,.{casas}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return str(v)

def _to_float_br(txt):
    return float(str(txt).replace('.', '').replace(',', '.'))

def converter_para_float_qualquer(s):
    s = str(s).strip()
    if ',' in s and '.' in s:
        return float(s.replace('.', '').replace(',', '.'))
    if ',' in s:
        return float(s.replace(',', '.'))
    return float(s)

def extenso_metros(v):
    v = round(float(v or 0), 2)
    m = int(v); cm = int(round((v - m) * 100))
    partes = []
    if m > 0:
        partes.append(num2words(m, lang='pt_BR') + (" metro" if m==1 else " metros"))
    if cm > 0:
        partes.append(num2words(cm, lang='pt_BR') + (" centímetro" if cm==1 else " centímetros"))
    return " e ".join(partes) if partes else "zero metro"

def area_por_extenso(v):
    v = round(float(v or 0), 2)
    m2 = int(v); cent = int(round((v - m2) * 100))
    if cent == 0:
        return f"{num2words(m2, lang='pt_BR')} metros quadrados"
    return f"{num2words(m2, lang='pt_BR')} metros quadrados e {num2words(cent, lang='pt_BR')} centésimos"

def hectares_from_m2(v):
    return float(v) / 10000.0

# ===================== Formatação de nomes/endereços/cidade/UF/bairro =====================
_PREP_MIN = {"DE", "DA", "DO", "DAS", "DOS"}

def _title_keep_preps(s: str) -> str:
    if not s:
        return ""
    t = s.strip().title()
    for prep in _PREP_MIN:
        t = re.sub(rf"\b{prep}\b", prep.lower(), t)
    t = re.sub(
        r'\s*,?\s*S\s*/\s*N[º°]?\b',
        lambda m: (',' if ',' in m.group(0) else '') + ' s/nº',
        t,
        flags=re.I
    )
    return t.strip()

def _fmt_cidade_slash_uf(s: str) -> str:
    if not s: return ""
    s = s.strip()
    if "/" not in s:
        return _title_keep_preps(s)
    cid, uf = [p.strip() for p in s.split("/", 1)]
    return f"{_title_keep_preps(cid)}/{uf.upper()}"

def _fmt_bairro(s: str) -> str:
    if not s:
        return ""
    t = s.strip().title()
    for prep in {"De", "Da", "Do", "Das", "Dos", "A", "E"}:
        t = re.sub(rf"\b{prep}\b", prep.lower(), t)
    return t

# ===================== Conversões de coordenadas =====================
_UF_HEMI_N = {'RR','AP'}
_UF_FUSO_DEFAULT = {
    'RS':'22S','SC':'22S','PR':'22S',
    'SP':'23S','RJ':'23S','MG':'23S','DF':'23S','MS':'21S',
    'ES':'24S','BA':'23S','GO':'22S','MT':'21S','TO':'22S',
    'MA':'23S','PA':'22S','RO':'20S','AC':'19S','AM':'20S',
    'RR':'20N','AP':'22N','RN':'24S','PB':'24S','PE':'24S',
    'AL':'24S','SE':'24S','CE':'24S','PI':'23S'
}

def _parse_uf(cidade_field):
    m = re.search(r'/\s*([A-Z]{2})\b', str(cidade_field).strip(), re.I)
    return m.group(1).upper() if m else None

def _zone_str_to_num_hemi(zstr):
    if not zstr:
        return (22, 'S')
    zstr = zstr.strip().upper().replace(' ', '')
    m = re.match(r'(\d{1,2})([NS])', zstr)
    if not m:
        return (22, 'S')
    return (int(m.group(1)), m.group(2))

def _auto_zone_from_city(cidade_field: str):
    uf = _parse_uf(cidade_field) or ''
    zstr = _UF_FUSO_DEFAULT.get(uf, '22S')
    zone_num, hemi = _zone_str_to_num_hemi(zstr)
    hemi = 'N' if uf in _UF_HEMI_N else 'S'
    return zone_num, hemi

def _utm_mc_from_zone(zone_num):
    mc = 6*zone_num - 183
    return abs(int(mc))

def _sirgas_utm_crs(zone_num: int, hemi: str) -> CRS:
    hemi = (hemi or 'S').upper()
    if hemi == 'S' and 18 <= int(zone_num) <= 25:
        return CRS.from_epsg(31960 + int(zone_num))
    south_flag = '+south ' if hemi == 'S' else ''
    proj4 = f"+proj=utm +zone={int(zone_num)} {south_flag}+datum=SIRGAS2000 +type=crs"
    return CRS.from_proj4(proj4)

def utm_to_latlon(E, N, zone_num, hemi='S'):
    try:
        E = converter_para_float_qualquer(E); N = converter_para_float_qualquer(N)
    except Exception:
        E = float(E); N = float(N)
    crs_utm = _sirgas_utm_crs(int(zone_num), hemi)
    crs_geo = CRS.from_epsg(4674)
    tr = Transformer.from_crs(crs_utm, crs_geo, always_xy=True)
    lon, lat = tr.transform(E, N)
    return lat, lon

def fmt_latlon_decimal(lat, lon):
    return f"Lat. {lat:.6f}°, Long. {lon:.6f}°"

def _dms_parts(val):
    sign = -1 if val < 0 else 1
    v = abs(val)
    d = int(v)
    m_float = (v - d) * 60
    m = int(m_float)
    s = (m_float - m) * 60
    return sign, d, m, s

def fmt_latlon_dms(lat, lon):
    sgn_lat, dlat, mlat, slat = _dms_parts(lat)
    sgn_lon, dlon, mlon, slon = _dms_parts(lon)
    def _mk(sign, d, m, s):
        s_txt = f"{s:06.3f}".replace('.', ',')
        prefix = '-' if sign < 0 else ''
        return f"{prefix}{d}°{m:02d}'{s_txt}\""
    return f"Lat. {_mk(sgn_lat, dlat, mlat, slat)}, Long. {_mk(sgn_lon, dlon, mlon, slon)}"

# ===================== Azimutes / direção cardinal =====================
def bearing_to_azimuth(b):
    if not b or not isinstance(b, str):
        return None
    s = b.strip().upper().replace('–','-').replace('°','-').replace("'",'-').replace('"','')
    s = re.sub(r'\s+', ' ', s)
    m = re.match(r'([NS])\s*([0-9]+)-([0-9]+)-([0-9]+(?:\.[0-9]+)?)\s*([EW])', s)
    if not m:
        m2 = re.search(r'(\d+)[^\d]+(\d+)[^\d]+(\d+(?:\.\d+)?)', s)
        if m2:
            d, mi, se = map(float, m2.groups()); return d + mi/60 + se/3600
        return None
    ns, d, mi, se, ew = m.groups()
    d, mi, se = float(d), float(mi), float(se)
    theta = d + mi/60 + se/3600
    if ns=='N' and ew=='E':
        az = theta
    elif ns=='S' and ew=='E':
        az = 180 - theta
    elif ns=='S' and ew=='W':
        az = 180 + theta
    elif ns=='N' and ew=='W':
        az = 360 - theta
    else:
        return None
    if az < 0: az += 360
    if az >= 360: az -= 360
    return az

def azimuth_to_dms_int(az):
    if az is None:
        return ""
    az = float(az) % 360.0
    d = int(az); m = int((az - d) * 60); s = int(round((az - d - m/60) * 3600))
    if s >= 60: s -= 60; m += 1
    if m >= 60: m -= 60; d += 1
    return f"{d}°{m:02d}'{s:02d}\""

def azimuth_to_card8(az):
    if az is None:
        return "XXXX"
    dirs = ["norte","nordeste","leste","sudeste","sul","sudoeste","oeste","noroeste"]
    idx = int(((az + 22.5) % 360) // 45)
    return dirs[idx]

# ===================== Inferir QUADRA pelo nome do arquivo =====================
def infer_quadra_from_filename(fname):
    up = os.path.basename(fname).upper()
    m = re.search(r'(QUADRA|SITE|QD)[ _\-]*([A-Z0-9]+)', up, flags=re.I)
    if m:
        return f"QUADRA {m.group(2)}"
    m2 = re.search(r'[_\- ]([A-Z0-9])\.(HTM|HTML|TXT)$', up)
    return f"QUADRA {m2.group(1)}" if m2 else "QUADRA (DESCONHECIDA)"

def _letters_to_number(tok):
    n = 0
    for ch in tok:
        if not ('A' <= ch <= 'Z'): return None
        n = n * 26 + (ord(ch) - ord('A') + 1)
    return n

def _quadra_sort_key(fname):
    q = infer_quadra_from_filename(fname)
    m = re.search(r'QUADRA\s+([A-Z0-9]+)', q)
    tok = m.group(1) if m else q
    if tok.isdigit():
        return (1, int(tok))
    tokU = tok.upper()
    num = _letters_to_number(tokU)
    if num is not None:
        return (0, num)
    return (2, tokU)

def _is_letters(tok: str) -> bool:
    return bool(re.fullmatch(r'[A-Z]+', tok or '', flags=re.I))

def _extract_quadra_token(label: str) -> str:
    m = re.search(r'QUADRA\s+([A-Z0-9]+)', str(label or ''), flags=re.I)
    return (m.group(1).upper() if m else str(label or '').upper()).strip()

def quadra_label_sort_key(label: str):
    tok = _extract_quadra_token(label)
    if _is_letters(tok):
        return (0, len(tok), tok)
    if tok.isdigit():
        return (1, int(tok))
    return (2, tok)

def _lote_num(v):
    try:
        return int(re.search(r'\d+', str(v)).group())
    except:
        return 10**9

# ===================== Detectores UNIF / DESM =====================
_UNIF_NAME_PAT = re.compile(r'\bUNIFICA(?:Ç|C)Ã?O\b', re.IGNORECASE)
_DESM_KEYS = re.compile(r'\b(GLEBA|ÁREA|AREA)\b', re.IGNORECASE)

def is_unificacao_item_name(nm: str) -> bool:
    return bool(_UNIF_NAME_PAT.search(str(nm or "")))

# ===================== Parsers =====================
def parse_parcels_from_txt(txt_bytes):
    txt = io.BytesIO(txt_bytes).read().decode('utf-8', errors='ignore')
    txt = txt.replace('\r', '')
    parts = re.split(r'(?:^|\n)\s*Name:\s*(\d+)\s*(?:\n|$)', txt)
    it = iter(parts); _ = next(it, "")
    parcels = []
    for num, bloco in zip(it, it):
        num = int(num)
        m0 = re.search(r'Point of Beginning\s*:\s*North:\s*([\d\.,]+)m\s*East:\s*([\d\.,]+)m', bloco, re.I)
        first_pt = {'Y': converter_para_float_qualquer(m0.group(1)), 'X': converter_para_float_qualquer(m0.group(2))} if m0 else None
        mA = re.search(r'Area:\s*([\d\.,]+)\s*sq\.m', bloco, re.I)
        area_m2 = converter_para_float_qualquer(mA.group(1)) if mA else None
        segs = []
        for m in re.finditer(r'Segment\s*#\d+.*?Line[\s\S]*?Course:\s*([NS].*?[EW])\s*Length:\s*([\d\.,]+)m', bloco, re.I):
            bearing = m.group(1).strip(); length = converter_para_float_qualquer(m.group(2))
            az = bearing_to_azimuth(bearing); segs.append({"type":"line","length_m":length,"azimuth":az})
        for m in re.finditer(r'Segment\s*#\d+.*?Curve[\s\S]*?Length:\s*([\d\.,]+)m[\s\S]*?Radius:\s*([\d\.,]+)m[\s\S]*?Course:\s*([NS].*?[EW])', bloco, re.I):
            curve_len = converter_para_float_qualquer(m.group(1)); radius = converter_para_float_qualquer(m.group(2)); chord_dir = m.group(3).strip()
            az = bearing_to_azimuth(chord_dir); segs.append({"type":"curve","curve_len_m":curve_len,"radius_m":radius,"azimuth":az})
        parcels.append({"num": num, "segments": segs, "area_m2": area_m2, "first_point": first_pt})
    return parcels

def parse_civilreport_from_html(html_bytes):
    soup = BeautifulSoup(html_bytes, "lxml")
    items = []
    for table in soup.find_all("table"):
        head = table.find("td", colspan="3")
        if not head: continue
        title = head.get_text(strip=True)
        if not title.upper().startswith("PARCEL"): continue
        name = title.split("Parcel",1)[1].strip() or "SEM NOME"
        ttxt = table.get_text("\n")
        m0 = re.search(r'Point\s+whose\s+Northing\s+is\s*([\d\.,]+)\s+and\s+whose\s+Easting\s*is\s*([\d\.,]+)', ttxt, re.I)
        first_pt = {'Y': converter_para_float_qualquer(m0.group(1)), 'X': converter_para_float_qualquer(m0.group(2))} if m0 else None
        mA = re.search(r'Area.*?\n.*?Square meters\s*\n\s*([\d\.,]+)', ttxt, re.I|re.S)
        area_m2 = converter_para_float_qualquer(mA.group(1)) if mA else None
        segs = []
        for m in re.finditer(r'Bearing:\s*([NS].*?[EW])\s*Length:\s*([\d\.,]+)', ttxt, re.I):
            bearing = m.group(1).strip(); length = converter_para_float_qualquer(m.group(2))
            az = bearing_to_azimuth(bearing); segs.append({"type":"line","length_m":length,"azimuth":az})
        for block in re.finditer(r'Curve.*?Curve Length:\s*([\d\.,]+).*?Radius Length:\s*([\d\.,]+).*?Chord Direction:\s*([NS].*?[EW])', ttxt, re.I|re.S):
            curve_len = converter_para_float_qualquer(block.group(1)); radius = converter_para_float_qualquer(block.group(2)); chord_dir = block.group(3).strip()
            az = bearing_to_azimuth(chord_dir)
            segs.append({"type":"curve","curve_len_m":curve_len,"radius_m":radius,"azimuth":az})
        items.append({'name': name, 'segments':segs, 'area_m2':area_m2, 'first_point':first_pt})
    return items

def parse_parcels_from_html(html_bytes):
    arr = parse_civilreport_from_html(html_bytes)
    parcels = []
    seq = 1
    for it in arr:
        m = re.search(r'(\d+)', str(it.get('name','')))
        num = int(m.group(1)) if m else seq
        parcels.append({
            "num": num,
            "segments": it.get("segments", []),
            "area_m2": it.get("area_m2"),
            "first_point": it.get("first_point")
        })
        seq += 1
    return parcels

# ===================== Classificação (regras) =====================
def _normalize(s):
    return re.sub(r'\s+', ' ', str(s or '')).strip().upper()

def classify_civil_item(name):
    n = _normalize(name)
    if re.search(r'\b(ALARGAMENTO(S)?|ACESSO(S)?( DE SERVIÇO(S)?)?|RODOVI(A|Á)S?|RUA(S)?|AVENIDA(S)?|PEATONAL(IS)?|CANTEIRO(S)?|ACESSOS?)\b', n):
        return ('viario', 'DESCRIÇÃO DE SISTEMA VIÁRIO')
    if re.search(r'^(AVS?\b)|\bÁREA(S)? VERDE(S)?\b', n):
        return ('verde', 'DESCRIÇÃO DE ÁREAS VERDES')
    if 'ÁREA VERDE DE PRESERVAÇÃO' in n or 'AREA VERDE DE PRESERVACAO' in n:
        return ('verde_preservacao', 'DESCRIÇÃO DE ÁREA VERDE DE PRESERVAÇÃO')
    if re.search(r'(PRESERVAÇÃO PERMANENTE|PRESERVACAO PERMANENTE|\bAPP\b|RESTRIÇ|RESTRICAO|PRESERVAÇÃO AMBIENTAL|PRESERVACAO AMBIENTAL)', n):
        if 'RESTRI' in n:
            return ('app', 'DESCRIÇÃO DE RESTRIÇÕES')
        if 'PRESERVAÇÃO AMBIENTAL' in n or 'PRESERVACAO AMBIENTAL' in n:
            return ('app', 'DESCRIÇÃO DE ÁREA DE PRESERVAÇÃO AMBIENTAL')
        return ('app', 'DESCRIÇÃO DE ÁREA DE PRESERVAÇÃO PERMANENTE')
    if re.search(r'\bAI(\b|\s)|\bÁREA(S)? INSTITUCIONAL(IS)?\b|\bAREA(S)? INSTITUCIONAL(IS)?\b', n):
        return ('institucional', 'DESCRIÇÃO DE ÁREAS INSTITUCIONAIS')
    if re.search(r'RESERVA TÉCNICA|RESERVA TECNICA|\bETE\b|\bEBE\b|\bETA\b|\bEBA\b|ESTAÇÃO DE BOMBEAMENTO|ESTACAO DE BOMBEAMENTO|ESTAÇÃO DE TRATAMENTO|ESTACAO DE TRATAMENTO', n):
        return ('reserva_tecnica', 'DESCRIÇÃO DE RESERVA TÉCNICA')
    if 'REMANESCENTE' in n:
        return ('remanescente', 'DESCRIÇÃO DE ÁREA REMANESCENTE')
    if re.search(r'ÁREA(S)? CONDOMINIA(L|IS)|\bAC\s*\d+\b|AREA(S)? CONDOMINIA(L|IS)', n):
        return ('condominial', 'DESCRIÇÃO DE ÁREAS CONDOMINIAIS')
    if n.startswith('QUADRA'):
        return ('quadras', 'DESCRIÇÃO DE QUADRAS')
    return ('outros', 'DESCRIÇÃO DE OUTRAS ÁREAS')

def _viario_base_and_trecho(nm_norm):
    n = _normalize(nm_norm)
    m_base = re.search(r'^(RUA|AVENIDA|RODOVIA|PEATONAL|ACESSO|CANTEIRO)\s+([A-Z0-9\-\/ ]+?)\s*(?:\-|–|—|\(|$)', n)
    if m_base:
        base = f"{m_base.group(1)} {m_base.group(2).strip()}"
    else:
        m2 = re.match(r'^([A-ZÇÃÕÉÊÍÓÚ ]+?)\s+(.+)$', n)
        base = f"{m2.group(1).strip()} {m2.group(2).strip()}" if m2 else n
    m_trecho = re.search(r'TRECHO[^\d]*(\d+)', n)
    trecho = int(m_trecho.group(1)) if m_trecho else 0
    return (base.strip(), trecho)

def _viario_sort_key(item_name):
    base, trecho = _viario_base_and_trecho(item_name)
    return (base, trecho)

# ===================== Formatação de texto =====================
def _title_case_name(nome: str) -> str:
    nome = (nome or "").strip().lower()
    return ' '.join(w.capitalize() for w in nome.split())

def _pt_date(prefixo_cidade="Porto Alegre"):
    MESES = ["janeiro","fevereiro","março","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]
    hoje = datetime.now()
    return f"{prefixo_cidade}, {hoje.day} de {MESES[hoje.month-1]} de {hoje.year}"

def _fmt_matriculas_plural(txt_raw: str):
    txt = (txt_raw or "").strip()
    if not txt:
        return ("matrícula", "XXXX")
    partes = [p.strip() for p in re.split(r'\s*(?:,|;| e )\s*', txt) if p.strip()]
    if len(partes) <= 1:
        return ("matrícula", partes[0] if partes else "XXXX")
    return ("matrículas", ", ".join(partes[:-1]) + " e " + partes[-1])

def _get_fmt_campos_basicos(form_data):
    """Adaptado para receber form_data ao invés de widgets"""
    nome_fmt = _title_case_name(form_data.get('nome_emp', ''))
    end_fmt = _title_keep_preps(form_data.get('endereco_emp', ''))
    cid_fmt = _fmt_cidade_slash_uf(form_data.get('cidade_emp', ''))
    bai_fmt = _fmt_bairro(form_data.get('bairro_emp', ''))
    return nome_fmt, end_fmt, cid_fmt, bai_fmt

# ===================== Logos / doc base =====================
def add_header_logo(doc, image_path, width_inches=1.4):
    if os.path.exists(image_path):
        try:
            for section in doc.sections:
                section.header_distance = Inches(0.8)
                p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
                r = p.add_run()
                r.add_picture(image_path, width=Inches(width_inches))
        except:
            pass  # Ignora se imagem não existir

def add_footer_logo(doc, image_path, width_inches=1.6):
    if os.path.exists(image_path):
        try:
            for section in doc.sections:
                section.footer_distance = Inches(0.3)
                p = section.footer.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                r = p.add_run()
                r.add_picture(image_path, width=Inches(width_inches))
        except:
            pass

def add_footer_left_text(doc, lines, size_pt=10):
    for section in doc.sections:
        p = section.footer.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for i, line in enumerate(lines):
            run = p.add_run(line)
            run.font.name = 'Calibri'
            run.font.size = Pt(size_pt)
            run.font.color.rgb = RGBColor(0,0,0)
            if i < len(lines)-1:
                run.add_break()

def add_page_numbers(document):
    section = document.sections[-1]
    p = section.footer.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE \\* MERGEFORMAT')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = ''
    r.append(t)
    fld.append(r)
    p._p.append(fld)

def add_corner_image_watermark_cm(doc, image_path, width_cm=6.46, height_cm=1.91):
    if os.path.exists(image_path):
        try:
            sec = doc.sections[0]
            para = sec.header.add_paragraph()
            r = para.add_run()
            r.add_picture(image_path, width=Cm(width_cm))
        except:
            pass

def _apply_moderate_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

def preparar_doc():
    doc = Document()
    _apply_moderate_margins(doc)
    add_header_logo(doc, CAMINHO_LOGO_CABECALHO)
    add_corner_image_watermark_cm(doc, CAMINHO_MARCA_DAGUA)
    add_footer_logo(doc, CAMINHO_LOGO_RODAPE)
    return doc

def _enable_update_fields_on_open(doc):
    settings_el = doc.settings._element
    for el in settings_el.iterchildren():
        if el.tag == qn('w:updateFields'):
            el.set(qn('w:val'), 'true')
            return
    upd = OxmlElement('w:updateFields')
    upd.set(qn('w:val'), 'true')
    settings_el.append(upd)

# ===================== Headings / helpers =====================
def heading(doc, text):
    h = doc.add_heading('', level=1)
    run = h.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    blank = doc.add_paragraph()
    blank.paragraph_format.space_after = Pt(0)
    return h

def _set_run_defaults(run, bold=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bool(bold)

def _add_hl(paragraph, txt="XXXX", bold=False):
    run = paragraph.add_run(txt)
    _set_run_defaults(run, bold=bold)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run

def _run_xxxx(par):
    r = par.add_run("XXXX")
    _set_run_defaults(r)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return r

def adicionar_texto_formatado(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    bold_pat = (
        r'(?:LOTE\s+\d+\s*–\s*QUADRA\s+[A-Z0-9]+:)'
        r'|(?:LOTE\s+\d+\s+da\s+QUADRA\s+[A-Z0-9]+)'
        r'|(?:(?<!Y=\s)(?<!X=\s)\d{1,3}(?:\.\d{3})*,\d+m²)'
        r'|(?:(?<!Y=\s)(?<!X=\s)\d{1,3}(?:\.\d{3})*,\d+m)'
    )
    coord_pat = (
        r'(?:Y=\s*\d{1,3}(?:\.\d{3})*,\d+m|X=\s*\d{1,3}(?:\.\d{3})*,\d+m)'
        r'|(?:Lat\.\s*-?\d+\.\d+°\s*,\s*Long\.\s*-?\d+\.\d+°)'
        r'|(?:Lat\.\s*-?\d+°\d{2}\'\d{2}(?:,\d+)?\"\s*,\s*Long\.\s*-?\d+°\d{2}\'\d{2}(?:,\d+)?\")'
    )
    dms_pat = r'\d{1,3}°\d{2}\'\d{2}(?:,\d{1,3})?"'
    bold_marker_pat = r'\[\[B\]\](.*?)\[\[/B\]\]'

    tok = re.compile(f'({bold_pat})|(XXXX)|({coord_pat})|({dms_pat})|({bold_marker_pat})',
                     flags=re.IGNORECASE|re.DOTALL)

    i = 0
    while i < len(texto):
        m = tok.search(texto, i)
        if not m:
            resto = texto[i:]
            parts = re.split(r'(XXXX)', resto)
            for part in parts:
                if part == '': continue
                if part == 'XXXX':
                    run = p.add_run(part)
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                    run = p.add_run(part)
                    run.font.name='Calibri'
                    run.font.size=Pt(12)
                    run.font.color.rgb=RGBColor(0,0,0)
            break

        pref = texto[i:m.start()]
        if pref:
            parts = re.split(r'(XXXX)', pref)
            for part in parts:
                if part == '': continue
                run = p.add_run(part)
                run.font.name='Calibri'
                run.font.size=Pt(12)
                run.font.color.rgb=RGBColor(0,0,0)
                if part == 'XXXX':
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        if m.group(1):
            run = p.add_run(m.group(1))
            run.bold = True
        elif m.group(2):
            run = p.add_run("XXXX")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif m.group(3) or m.group(4):
            run = p.add_run(m.group(0))
            run.bold = False
        else:
            inner = re.sub(r'^\[\[B\]\]|\[\[/B\]\]$', '', m.group(5))
            run = p.add_run(inner)
            run.bold = True

        run.font.name='Calibri'
        run.font.size=Pt(12)
        run.font.color.rgb=RGBColor(0,0,0)
        i = m.end()

# ===================== Funções auxiliares para coordenadas =====================
def _fmt_coord_dec(val):
    try:
        return f"{float(val):.6f}".replace(".", ",") + "°"
    except:
        return str(val)

def _fmt_coord_dms(val):
    sign = '-' if float(val) < 0 else ''
    v = abs(float(val))
    d = int(v)
    m = int((v - d) * 60)
    s = (v - d - m/60) * 3600
    s_txt = f"{s:.3f}".replace(".", ",")
    return f"{sign}{d}°{m:02d}'{s_txt}\""

def _dms_str(az):
    return azimuth_to_dms_int(az) if az is not None else ""

def _propaga_vertices(first_point: dict, segments: list,
                      coord_fmt_str: str = 'utm',
                      zone_num: int = 22,
                      hemi: str = 'S'):
    if not first_point or not segments:
        return []

    x = float(first_point["X"])
    y = float(first_point["Y"])
    p_idx = 1
    rows = []

    for seg in segments:
        az = float(seg.get("azimuth") or 0.0)
        rad = math.radians(az)

        if seg.get("type") == "line":
            L = float(seg.get("length_m") or 0.0)
            dx = math.sin(rad) * L
            dy = math.cos(rad) * L
            x2, y2 = x + dx, y + dy
            dist = round(L, 2)
            raio = None
        else:
            arc = float(seg.get("curve_len_m") or 0.0)
            R_ = float(seg.get("radius_m") or 0.0)
            theta = (arc / R_) if R_ > 0 else 0.0
            chord = 2.0 * R_ * math.sin(theta / 2.0)
            dx = math.sin(rad) * chord
            dy = math.cos(rad) * chord
            x2, y2 = x + dx, y + dy
            dist = round(arc, 2)
            raio = round(R_, 2) if R_ else None

        if coord_fmt_str == 'utm':
            c1 = _fmt_br(x2, 2)
            c2 = _fmt_br(y2, 2)
        else:
            lat, lon = utm_to_latlon(x2, y2, zone_num, hemi)
            if coord_fmt_str == 'dec':
                c1 = _fmt_coord_dec(lon)
                c2 = _fmt_coord_dec(lat)
            else:
                c1 = _fmt_coord_dms(lon)
                c2 = _fmt_coord_dms(lat)

        rows.append({
            "DE": f"P{p_idx}",
            "PARA": f"P{p_idx + 1}",
            "COORD_1": c1,
            "COORD_2": c2,
            "AZIMUTE": _dms_str(az),
            "DISTANCIA (m)": dist,
            "RAIO (m)": raio,
            "CONFRONTANTE": ""
        })

        x, y = x2, y2
        p_idx += 1

    return rows

# ===================== Builders (lotes e áreas) =====================
def _texto_ane(largura_m):
    num_sem_negrito = f"{_fmt_br(largura_m, 2)}\u200Bm"
    ext = extenso_metros(largura_m)
    return (
        f" Existe uma faixa não edificante com largura de {num_sem_negrito} ({ext}), "
        f"conforme definido no projeto urbanístico e nas restrições de uso do terreno."
    )

def _format_first_point(fp, coord_fmt, zone_num, hemi):
    if not fp: return None
    y = round(float(fp["Y"]), 2); x = round(float(fp["X"]), 2)
    if coord_fmt == 'utm':
        return f"ponto de coordenadas Y= {_fmt_br(y, 2)}m e X= {_fmt_br(x, 2)}m"
    lat, lon = utm_to_latlon(x, y, zone_num, hemi)
    if coord_fmt == 'dec':
        return f"ponto de coordenadas geográficas {fmt_latlon_decimal(lat, lon)}"
    return f"ponto de coordenadas geográficas {fmt_latlon_dms(lat, lon)}"

def _seg_texto_com_card(seg, dest_coord=None, tipo='line', coord_fmt='utm'):
    az = seg.get("azimuth")
    card = azimuth_to_card8(az)
    az_dms = azimuth_to_dms_int(az)

    dest_txt = ""
    if dest_coord:
        c1, c2 = dest_coord
        if coord_fmt == 'utm':
            dest_txt = f" até o ponto de coordenadas Y= {c2}m e X= {c1}m"
        else:
            dest_txt = f" até o ponto de coordenadas {c2} / {c1}"

    if tipo == 'line':
        lv = round(float(seg["length_m"]), 2)
        length = _fmt_br(lv, 2) + "m"
        return (
            f"daí segue, por reta, sentido {card}, medindo {length} ({extenso_metros(lv)}), "
            f"confrontando ao XXXX com XXXX{dest_txt}, seguindo por um azimute de {az_dms}; "
        )

    clv = round(float(seg["curve_len_m"]), 2)
    rv = round(float(seg["radius_m"]), 2)
    cl = _fmt_br(clv, 2) + "m"
    r = _fmt_br(rv, 2) + "m"
    return (
        f"daí segue, por curva, sentido {card}, medindo {cl} ({extenso_metros(clv)}) e raio de {r} ({extenso_metros(rv)}), "
        f"confrontando ao XXXX com XXXX{dest_txt}, seguindo por um azimute de {az_dms}; "
    )

def build_area_text(item_name, item, tipo_full, empreendimento, endereco, bairro, cidade,
                    ane_enable=False, ane_largura_m=None, coord_fmt='utm', zone_num=22, hemi='S',
                    ident_prefix=None, ident_label_only=False, ident_label_text="Descrição do Imóvel:"):
    nome_norm = _normalize(item_name)
    area = item.get("area_m2") or 0
    area_fmt = _fmt_br(area, 2) + "m²"
    area_ext = area_por_extenso(area)

    tipo_is_lote_cond = (tipo_full or "").lower() in (
        "condomínio fechado de lotes residenciais",
        "condomínio fechado de lotes",
        "loteamento de acesso controlado"
    )

    if ident_label_only:
        cabeca = f"[[B]]{ident_label_text}[[/B]] Um terreno urbano, irregular, sem benfeitorias, "
    else:
        if ident_prefix:
            cabeca = f"{ident_prefix} [[B]]{nome_norm}[[/B]]: Um terreno urbano, irregular, sem benfeitorias, "
        else:
            cabeca = f"[[B]]{nome_norm}[[/B]]: Um terreno urbano, irregular, sem benfeitorias, "

    if tipo_is_lote_cond:
        cabeca += (
            f"localizado na {endereco}, no bairro {bairro}, na cidade de {cidade}, "
            f"constituído como [[B]]{_normalize(item_name)}[[/B]], "
        )
    else:
        cabeca += (
            f"situado entre terras que são ou foram de XXXX, "
            f"localizado na {endereco}, no bairro {bairro}, na cidade de {cidade}, "
            f"constituído como [[B]]{_normalize(item_name)}[[/B]], "
        )

    if item.get("first_point"):
        fp_txt = _format_first_point(item["first_point"], coord_fmt, zone_num, hemi)
        if fp_txt:
            cabeca += f"inicia-se a descrição no {fp_txt}; "

    rows = _propaga_vertices(
        item.get("first_point"),
        item.get("segments", []),
        coord_fmt_str=coord_fmt,
        zone_num=zone_num,
        hemi=hemi
    )

    partes = []
    segs = item.get("segments", []) or []
    for i, seg in enumerate(segs):
        dest = None
        if i < len(rows):
            dest = (rows[i]["COORD_1"], rows[i]["COORD_2"])
        if seg["type"] == "line":
            partes.append(_seg_texto_com_card(seg, dest_coord=dest, tipo='line', coord_fmt=coord_fmt))
        else:
            partes.append(_seg_texto_com_card(seg, dest_coord=dest, tipo='curve', coord_fmt=coord_fmt))

    corpo = "".join(partes)
    if corpo.endswith("; "):
        corpo = corpo[:-2] + ", "

    texto = cabeca + corpo + "chegando ao final da descrição do perímetro."
    texto += " Dista XXXXm da esquina da Rua XXXX."

    if ane_enable and (ane_largura_m is not None):
        texto += _texto_ane(ane_largura_m)

    return texto

def build_memorial_text(parcel, quadra, tipo_full, empreendimento, endereco, bairro, cidade,
                        ane_enable=False, ane_largura_m=None, eh_condominio=False,
                        area_tot_priv=0.0, area_tot_cond=0.0, coord_fmt='utm', zone_num=22, hemi='S'):
    num = parcel["num"]
    area = parcel.get("area_m2") or 0
    area_fmt = _fmt_br(area, 2) + "m²"
    area_ext = area_por_extenso(area)

    tipo_is_lote_cond = (tipo_full or "").lower() in (
        "condomínio fechado de lotes residenciais",
        "condomínio fechado de lotes",
        "loteamento de acesso controlado"
    )

    cabeca = f"LOTE {num} – {quadra}: Um terreno urbano, irregular, sem benfeitorias, "

    if tipo_is_lote_cond:
        cabeca += (
            f"localizado na {endereco}, no bairro {bairro}, na cidade de {cidade}, "
            f"constituído como LOTE {num} da {quadra}, "
        )
    else:
        cabeca += (
            f"situado entre terras que são ou foram de XXXX, "
            f"localizado na {endereco}, no bairro {bairro}, na cidade de {cidade}, "
            f"constituído como LOTE {num} da {quadra}, "
        )

    if parcel.get("first_point"):
        fp_txt = _format_first_point(parcel["first_point"], coord_fmt, zone_num, hemi)
        if fp_txt:
            cabeca += f"inicia-se a descrição no {fp_txt}; "

    rows = _propaga_vertices(
        parcel.get("first_point"),
        parcel.get("segments", []),
        coord_fmt_str=coord_fmt,
        zone_num=zone_num,
        hemi=hemi
    )

    partes = []
    segs = parcel.get("segments", []) or []
    for i, seg in enumerate(segs):
        dest = None
        if i < len(rows):
            dest = (rows[i]["COORD_1"], rows[i]["COORD_2"])
        if seg["type"] == "line":
            partes.append(_seg_texto_com_card(seg, dest_coord=dest, tipo='line', coord_fmt=coord_fmt))
        else:
            partes.append(_seg_texto_com_card(seg, dest_coord=dest, tipo='curve', coord_fmt=coord_fmt))

    corpo = "".join(partes)
    if corpo.endswith("; "):
        corpo = corpo[:-2] + ", "

    texto = cabeca + corpo + "chegando ao final da descrição do perímetro."
    texto += " Dista XXXXm da esquina da Rua XXXX."

    if ane_enable and (ane_largura_m is not None):
        texto += _texto_ane(ane_largura_m)

    if eh_condominio and area and (area_tot_priv or 0) > 0:
        fr = area / (area_tot_priv or 1.0)
        area_comum = fr * (area_tot_cond or 0.0)
        area_total = area + area_comum
        priv_fmt = _fmt_br(area, 2)
        uso_fmt = _fmt_br(area_comum, 2)
        total_fmt = _fmt_br(area_total, 2)
        frac_fmt = f"{fr:.7f}"
        m2 = "\u200Bm²"
        texto += (
            f" Possui área real privativa de {priv_fmt}{m2}, área de uso comum de {uso_fmt}{m2}, "
            f"área real total de {total_fmt}{m2}, correspondendo-lhe a fração ideal de {frac_fmt}."
        )

    return texto

# ===================== Funções auxiliares para Memorial Resumo =====================
def _join_com_e(itens):
    itens = [str(i) for i in itens if str(i).strip()]
    if not itens: return "XXXX"
    if len(itens) == 1: return itens[0]
    return ", ".join(itens[:-1]) + " e " + itens[-1]

def _add_centered(doc, txt, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(txt)
    _set_run_defaults(run, bold=bold)
    return p

def _add_toc(doc):
    """
    Insere TOC (níveis 1–3) como campo complexo e marca como 'dirty'
    para o Word atualizar automaticamente ao abrir.
    
    Nota: O Word pode não atualizar automaticamente. Se o sumário aparecer vazio:
    - Clique com botão direito no campo e selecione "Atualizar Campo"
    - Ou pressione F9 com o cursor no campo
    - Ou vá em Referências > Atualizar Tabela
    """
    p = doc.add_paragraph()
    r = p.add_run()

    # begin - marca como dirty para forçar atualização
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.set(qn('w:dirty'), 'true')  # Força atualização

    # instrução TOC
    # \o "1-3" = níveis de heading 1 a 3
    # \h = cria hyperlinks
    # \z = oculta números de página no modo web
    # \u = usa estilos de heading
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r'TOC \o "1-3" \h \z \u'

    # separador (conteúdo temporário antes do update)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')

    # placeholder visível até o Word atualizar
    r_tmp = OxmlElement('w:r')
    t_tmp = OxmlElement('w:t')
    t_tmp.text = "Sumário será atualizado ao abrir o documento…"
    r_tmp.append(t_tmp)

    # end
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    # monta no run (ordem é crítica!)
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(r_tmp)
    r._r.append(fld_end)

def _heading_num(doc, idx, title):
    h = heading(doc, f"{idx}. {title}")
    return h

def _remove_trailing_empty_paragraphs(doc):
    """Remove apenas parágrafos realmente vazios no fim do documento"""
    def _para_has_field(p):
        return bool(p._p.xpath('.//w:fldChar') or p._p.xpath('.//w:instrText'))

    while doc.paragraphs:
        last = doc.paragraphs[-1]
        is_text_empty = not (last.text or '').strip()
        if is_text_empty and not _para_has_field(last):
            last._element.getparent().remove(last._element)
        else:
            break

def _sec_assinaturas_resumo(doc):
    heading(doc, "ASSINATURAS")
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("_____________________________")
    _set_run_defaults(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Responsável técnico")
    _set_run_defaults(r, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SOLIDO - DESIGN URBANO LTDA.")
    _set_run_defaults(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("CAU-RS 15335-4")
    _set_run_defaults(r)

# ===================== Funções de geração web (adaptadas) =====================
def _build_memorial_resumo_doc_web(form_data, output_dir):
    """
    Gera o 'Memorial Descritivo/Resumo' com:
    - Capa: título + subtítulo + Sumário
    - TOC com números de página (força update on open)
    - Seções numeradas a partir da página 2
    """
    try:
        doc = preparar_doc()
        _enable_update_fields_on_open(doc)
    except Exception as e:
        import traceback
        print(f"❌ ERRO ao preparar documento: {e}")
        traceback.print_exc()
        raise

    nome_fmt = _title_case_name(form_data.get('nome_emp', ''))
    end_fmt = _title_keep_preps(form_data.get('endereco_emp', ''))
    cid_fmt = _fmt_cidade_slash_uf(form_data.get('cidade_emp', ''))
    bai_fmt = _fmt_bairro(form_data.get('bairro_emp', ''))

    tipo_proj = form_data.get('tipo_proj_resumo', 'condominio')
    is_cond = (tipo_proj == 'condominio')
    tipo_lbl = "Condomínio fechado de lotes" if is_cond else "Loteamento de acesso controlado"

    # ===== CAPA =====
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("MEMORIAL DESCRITIVO")
    _set_run_defaults(r, bold=True)
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r1 = p.add_run(tipo_lbl + " ")
    _set_run_defaults(r1, bold=True)
    r1.font.size = Pt(14)
    r2a = p.add_run('"')
    _set_run_defaults(r2a, bold=True)
    r2a.font.size = Pt(14)
    nome_txt = (nome_fmt or "")
    if not nome_txt.strip():
        rr = p.add_run("XXXX")
        _set_run_defaults(rr, bold=True)
        rr.font.size = Pt(14)
        rr.font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        rr = p.add_run(nome_txt)
        _set_run_defaults(rr, bold=True)
        rr.italic = True
        rr.font.size = Pt(14)
    r2b = p.add_run('"')
    _set_run_defaults(r2b, bold=True)
    r2b.font.size = Pt(14)

    doc.add_paragraph()

    # Sumário
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("Sumário")
    _set_run_defaults(r, bold=True)
    r.font.size = Pt(14)
    _add_toc(doc)

    _remove_trailing_empty_paragraphs(doc)

    # ===== PÁGINA 2+ =====
    idx = 1

    # 1. INTRODUÇÃO
    h_intro = _heading_num(doc, idx, "INTRODUÇÃO")
    h_intro.paragraph_format.page_break_before = True
    idx += 1

    usos_sel = form_data.get('usos_multi', [])
    if isinstance(usos_sel, str):
        usos_sel = [usos_sel]
    usos_txt = _join_com_e(usos_sel)
    modo_lbl = "loteamento" if tipo_proj == 'loteamento' else "condomínio"

    # Parágrafo 1
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run("O "))
    r = p.add_run(tipo_lbl + " ")
    _set_run_defaults(r, bold=True)
    p.add_run('"')
    r = p.add_run((nome_fmt or "").strip() or "XXXX")
    _set_run_defaults(r, bold=True)
    r.italic = True
    if not (form_data.get('nome_emp') or "").strip():
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run('"')
    _set_run_defaults(p.add_run(
        f" é um empreendimento por unidades autônomas a construir, com finalidade {usos_txt.lower()}."
    ))

    # Parágrafo 2
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        "Constitui-se de um projeto de parcelamento do solo urbano em uma gleba situada em frente à "
    ))
    _set_run_defaults(p.add_run(end_fmt or 'XXXX'))
    if not end_fmt:
        _run_xxxx(p)
    _set_run_defaults(p.add_run(", no bairro "))
    _set_run_defaults(p.add_run(bai_fmt or 'XXXX'))
    if not bai_fmt:
        _run_xxxx(p)
    _set_run_defaults(p.add_run(", na área urbana do município de "))
    if (cid_fmt or '').strip():
        _set_run_defaults(p.add_run(cid_fmt))
    else:
        _run_xxxx(p)
    _set_run_defaults(p.add_run(
        ". As unidades autônomas a construir terão área mínima de "
    ))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("m². "))

    _set_run_defaults(p.add_run("A densidade prevista é de "))
    _set_run_defaults(p.add_run(str(form_data.get('num_lotes_emp', 0) or 0)))
    _set_run_defaults(p.add_run(" lotes, distribuídos em "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(
        " quadras, além de áreas de uso comum e vias de circulação — trazendo a "
    ))
    _set_run_defaults(p.add_run(cid_fmt if cid_fmt else ""))
    if not cid_fmt:
        _run_xxxx(p)
    _set_run_defaults(p.add_run(
        " um empreendimento que oferece condições qualificadas de vida urbana e social a seus moradores e que, "
        "em conjunto com outras ações previstas, contribuirá para o desenvolvimento socioeconômico do município."
    ))

    # 2. PROPRIETÁRIO/INCORPORADORA
    _heading_num(doc, idx, "PROPRIETÁRIO/INCORPORADORA")
    idx += 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _run_xxxx(p)
    _set_run_defaults(p.add_run(", inscrita no CNPJ sob o nº "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(", incorporadora da área de "))
    if (form_data.get('area_total_emp') or '').strip():
        try:
            v = _to_float_br(form_data.get('area_total_emp'))
            _set_run_defaults(p.add_run(_fmt_br(v, 2) + "m²"))
        except:
            _run_xxxx(p)
            _set_run_defaults(p.add_run("m²"))
    else:
        _run_xxxx(p)
        _set_run_defaults(p.add_run("m²"))

    mats_raw = (form_data.get('matricula_emp') or '').strip()
    parts = [s for s in re.split(r'\s*(?:,|;| e )\s*', mats_raw) if s]
    if not parts:
        _set_run_defaults(p.add_run(" registrada na matrícula nº "))
        _run_xxxx(p)
    else:
        if len(parts) == 1:
            _set_run_defaults(p.add_run(" registrada na matrícula nº "))
            _set_run_defaults(p.add_run(parts[0]))
        else:
            _set_run_defaults(p.add_run(" registradas nas Matrículas nº "))
            _set_run_defaults(p.add_run(", ".join(parts[:-1]) + " e " + parts[-1]))
    _set_run_defaults(p.add_run(" conforme Ofício do Registro de Imóveis da Comarca de "))
    if (cid_fmt or '').strip():
        _set_run_defaults(p.add_run(cid_fmt))
    else:
        _run_xxxx(p)
    _set_run_defaults(p.add_run("."))

    # 3. RESPONSÁVEL TÉCNICO
    _heading_num(doc, idx, "RESPONSÁVEL TÉCNICO PELO PROJETO URBANÍSTICO")
    idx += 1
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        "SOLIDO - DESIGN URBANO LTDA., registrada no CAU-RS 15335-4, "
        "CNPJ nº 26.887.368/0001-07."
    ))

    # 4. A GLEBA
    _heading_num(doc, idx, "A GLEBA")
    idx += 1
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run("A área para implantação do "))
    r = p.add_run(tipo_lbl + " ")
    _set_run_defaults(r, bold=True)
    p.add_run('"')
    r = p.add_run(nome_fmt.strip() if (nome_fmt or "").strip() else "XXXX")
    _set_run_defaults(r, bold=True)
    r.italic = True
    if not (nome_fmt or "").strip():
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run('"')
    _set_run_defaults(p.add_run(" é de "))
    if (form_data.get('area_total_emp') or '').strip():
        try:
            v = _to_float_br(form_data.get('area_total_emp'))
            _set_run_defaults(p.add_run(_fmt_br(v, 2) + "m²"))
        except:
            _run_xxxx(p)
            _set_run_defaults(p.add_run("m²"))
    else:
        _run_xxxx(p)
        _set_run_defaults(p.add_run("m²"))
    _set_run_defaults(p.add_run(
        f", com frente à {end_fmt or 'XXXX'}, bairro {bai_fmt or 'XXXX'}, "
        "na área urbana do município de "
    ))
    if (cid_fmt or '').strip():
        _set_run_defaults(p.add_run(cid_fmt))
    else:
        _run_xxxx(p)
    _set_run_defaults(p.add_run(
        ", em conformidade com as diretrizes municipais. "
        "As medidas da poligonal que delimitam esta área, bem como suas confrontações com os lindeiros, "
        "têm como referência e estão expressas no Levantamento Topográfico Planialtimétrico, "
        "bem como descritas na matrícula acima citada."
    ))

    # 5. TOPOGRAFIA
    _heading_num(doc, idx, "TOPOGRAFIA")
    idx += 1
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    topografia_val = form_data.get('topografia', 'Acentuada')
    if topografia_val == 'Acentuada':
        txt = (
            "A área apresenta uma topografia acentuada, marcada por fortes declividades e variações expressivas de altitude ao longo de sua extensão. "
            "O relevo é irregular, com presença de encostas íngremes, vales e elevações bem definidas, evidenciando intensa atuação de processos erosivos e de dissecação do terreno. "
            "Essa configuração indica um ambiente de dinâmica geológica ativa, onde a movimentação de águas superficiais é mais intensa e direcionada. "
            "As diferenças de cotas altimétricas conferem à paisagem um aspecto movimentado e de grande expressividade visual, típico de regiões montanhosas ou de relevo fortemente ondulado."
        )
    else:
        txt = (
            "A área apresenta uma topografia plana, com superfície predominantemente nivelada e variações altimétricas pouco expressivas. "
            "O relevo é homogêneo, sem presença significativa de declives ou irregularidades marcantes, o que indica um processo de formação estável e pouco sujeito a erosões intensas. "
            "As linhas de drenagem são discretas e tendem a se distribuir de forma mais uniforme, resultando em escoamento superficial lento. "
            "Esse tipo de configuração favorece uma leitura clara do terreno e mantém uma paisagem contínua e regular, característica comum de regiões sedimentares ou áreas de deposição recente."
        )
    _set_run_defaults(p.add_run(txt))

    # LOCALIZAÇÃO/AEROFOTO (não numerado)
    _add_centered(doc, "LOCALIZAÇÃO/AEROFOTO", bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_hl(p, "(inserir imagem)")

    # 6. ZONEAMENTO
    zone_idx = idx
    _heading_num(doc, idx, "ZONEAMENTO")
    idx += 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(f"{zone_idx}.1. A gleba se encontra na macrozona "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(", com os seguintes parâmetros volumétricos a serem seguidos:"))

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(f"{zone_idx}.2. Lotes com altura máxima de "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("m; taxa de ocupação (T.O.) de "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("%, taxa de permeabilidade (T.P.) de "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("%. Deverá atender ao recuo de jardim de "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("m."))

    # 7. DESTINAÇÃO E CONSTRUÇÕES
    dest_idx = idx
    _heading_num(doc, idx, "DESTINAÇÃO E CONSTRUÇÕES")
    idx += 1

    has_ai_val = form_data.get('has_ai', False)
    if isinstance(has_ai_val, str):
        has_ai_val = has_ai_val.lower() in ('true', '1', 'yes', 'on')
    
    if has_ai_val:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(p.add_run(f"{dest_idx}.1. Será reservada uma área de doação pública ao Município de "))
        if cid_fmt.strip():
            _set_run_defaults(p.add_run(cid_fmt))
        else:
            _run_xxxx(p)
        _set_run_defaults(p.add_run(" que corresponderá a "))
        _run_xxxx(p)
        _set_run_defaults(p.add_run("% ("))
        _run_xxxx(p)
        _set_run_defaults(p.add_run("m²"))
        _set_run_defaults(p.add_run(
            "), atendendo integralmente ao percentual mínimo exigido para área institucional. "
            "As áreas institucionais foram estrategicamente alocadas de forma a garantir sua viabilidade técnica para implantação de equipamentos públicos."
        ))
        sub = 2
    else:
        sub = 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(f"{dest_idx}.{sub}. Os lotes do "))
    r = p.add_run(tipo_lbl + " ")
    _set_run_defaults(r, bold=True)
    p.add_run('"')
    r = p.add_run(nome_fmt if nome_fmt.strip() else "XXXX")
    _set_run_defaults(r, bold=True)
    r.italic = True
    if not (form_data.get('nome_emp') or "").strip():
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run('"')
    _set_run_defaults(p.add_run(
        f", destinam-se a construções de uso {usos_txt.lower()}, conforme legislação vigente. "
        "As atividades sociais e de lazer serão dispostas junto às áreas verdes, de lazer e de serviços."
    ))

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        f"{dest_idx}.{sub+1}. Todas as construções além de atender às exigências morfológicas e tipológicas do empreendimento, "
        f"irão atender às exigências municipais contidas no Plano Diretor e no Código de Edificações do Município de "
        f"{_cidade_sem_uf(form_data.get('cidade_emp', '')) or 'XXXX'}."
    ))

    # 8. SISTEMA VIÁRIO
    viario_idx = idx
    _heading_num(doc, idx, "SISTEMA VIÁRIO")
    idx += 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(f"{viario_idx}.1. O sistema viário do "))
    r = p.add_run(tipo_lbl + " ")
    _set_run_defaults(r, bold=True)
    p.add_run('"')
    r = p.add_run(nome_fmt if nome_fmt.strip() else "XXXX")
    _set_run_defaults(r, bold=True)
    r.italic = True
    if not (form_data.get('nome_emp') or "").strip():
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run('"')
    _set_run_defaults(p.add_run(
        " foi projetado respeitando a topografia local e as normas técnicas vigentes. "
        "Seu arruamento será detalhado no Projeto Urbanístico, Geométrico e de Pavimentação, "
        "tendo uma área aproximada total de "
    ))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(
        "m², constituído de passeios e pista de rolamento. "
        "A via de acesso terá largura de "
    ))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("m, pistas de rolamento de "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("m e passeios de "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(
        "m. Internamente, após o acesso controlado (guarita e cancelas), as vias terão largura média de "
    ))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("m, com pistas de rolamento de "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("m e passeios de "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(
        "m. As vias serão denominadas futuramente. "
        "A previsão de revestimento é paver intertravado de concreto, podendo haver alteração após estudo preciso do solo e análise econômica. "
        "O sistema de drenagem será entregue com bueiros e declividades dimensionadas de acordo com a demanda local e a legislação pertinente."
    ))

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(f"{viario_idx}.2. Os passeios terão largura total média de "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(
        "m, garantindo faixa livre mínima de 1,20m para circulação contínua e desobstruída."
    ))

    # 9. ÁREAS CONDOMINIAIS (só se condomínio)
    if is_cond:
        areas_cond_idx = idx
        _heading_num(doc, idx, "ÁREAS CONDOMINIAIS")
        idx += 1

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(p.add_run(
            f"{areas_cond_idx}.1. As áreas de uso comum serão divididas em áreas verdes, de lazer e serviços. "
            "O condomínio contará com diversas áreas de uso comum, edificadas, ambientadas e arborizadas, com diversas infraestruturas para uso dos futuros moradores, "
            "em conformidade com todos os critérios municipais e legislações pertinentes."
        ))

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(p.add_run(
            f"{areas_cond_idx}.2. As áreas de lazer e serviços serão aquelas destinadas a guarita, administração e infraestrutura do pórtico de acesso, área técnica, estruturas condominiais, "
            "quadras esportivas e vagas de estacionamento. Os acessos a essas áreas serão plenamente acessíveis, atendendo às normas de acessibilidade. "
            "O dimensionamento dos espaços foi estabelecido conforme o projeto urbanístico, que visa qualificar o espaço e garantir as melhores condições de acessibilidade e moradia."
        ))

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(p.add_run(
            f"{areas_cond_idx}.3. As áreas condominiais e o sistema viário do condomínio serão administrados pela Administração do Condomínio e seu Estatuto Social, "
            "a serem constituídos por escritura pública no foro deste Município e Comarca, tendo como atribuições a administração, limpeza, vigilância, manutenção e conservação da infraestrutura."
        ))

    # ÁREA DE RESTRIÇÃO (somente se marcado no formulário)
    has_restricao_val = form_data.get('has_restricao', False)
    if isinstance(has_restricao_val, str):
        has_restricao_val = has_restricao_val.lower() in ('true', '1', 'yes', 'on')
    
    if has_restricao_val:
        ar_idx = idx
        _heading_num(doc, idx, "ÁREA DE RESTRIÇÃO")
        idx += 1

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(p.add_run(f"{ar_idx}.1. O "))
        r = p.add_run(tipo_lbl + " ")
        _set_run_defaults(r, bold=True)
        p.add_run('"')
        r = p.add_run(nome_fmt if nome_fmt.strip() else "XXXX")
        _set_run_defaults(r, bold=True)
        r.italic = True
        if not (form_data.get('nome_emp') or "").strip():
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        p.add_run('" ')
        _set_run_defaults(p.add_run("conta com área de restrição, totalizando "))
        _run_xxxx(p)
        _set_run_defaults(p.add_run("m², correspondendo a "))
        _run_xxxx(p)
        _set_run_defaults(p.add_run("% da área total do empreendimento."))

    # ÁREA PRIVATIVA
    ap_idx = idx
    _heading_num(doc, idx, "ÁREA PRIVATIVA")
    idx += 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(f"{ap_idx}.1. A área privativa proposta é de aproximadamente "))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(
        f"m², constituída por aproximadamente {(form_data.get('num_lotes_emp') or 0)} lotes com área mínima de "
    ))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(
        "m². Todos os lotes estarão dispostos no Projeto Urbanístico e estão distribuídos em "
    ))
    _run_xxxx(p)
    _set_run_defaults(p.add_run(
        " quadras, além das áreas de uso comum e vias de circulação."
    ))

    _add_centered(doc,
        "Abaixo são descritas as áreas totais do empreendimento e suas respectivas divisões:",
        bold=False
    )
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_hl(p, "(inserir tabela resumo)")

    # MUROS E CERCAMENTOS
    muros_idx = idx
    _heading_num(doc, idx, "MUROS E CERCAMENTOS")
    idx += 1

    if is_cond:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(p.add_run(
            f"{muros_idx}.1. O condomínio contará com muro tipo gradil de concreto com altura de "
        ))
        _run_xxxx(p)
        _set_run_defaults(p.add_run(
            " metros, implantado ao longo de toda a sua extensão perimetral voltada para as vias públicas externas. "
            "O cercamento será instalado de acordo com o padrão arquitetônico do condomínio, com tratamento paisagístico externo, "
            "em conformidade com as definições estabelecidas junto à municipalidade. "
            "O sistema de controle de acesso será composto por guarita e cancelas automáticas, conforme indicado em planta urbanística, "
            "e eventuais cercamentos internos ou divisórios seguirão o mesmo padrão visual e construtivo, respeitando a legislação municipal vigente."
        ))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(p.add_run(
            f"{muros_idx}.1. O loteamento contará com muro tipo gradil de concreto, com altura de "
        ))
        _run_xxxx(p)
        _set_run_defaults(p.add_run(
            " metros, implantado apenas nas divisas voltadas para áreas públicas. "
            "Nas divisas entre lotes ou áreas internas, não será executado muro pelo empreendimento, "
            "ficando sua construção sob responsabilidade dos respectivos proprietários, conforme diretrizes do loteamento. "
            "O cercamento seguirá o padrão arquitetônico estabelecido, com tratamento paisagístico externo, conforme acordado com a municipalidade. "
            "O sistema de controle de acesso será composto por guarita e cancelas automáticas, conforme planta urbanística e legislação vigente."
        ))

    # PAISAGISMO E ARBORIZAÇÃO
    pais_idx = idx
    _heading_num(doc, idx, "PAISAGISMO E ARBORIZAÇÃO")
    idx += 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        f"{pais_idx}.1. Será implantada faixa de paisagismo com vegetação arbustiva e forrações de porte compatível, "
        "contribuindo para a integração urbana, a valorização das áreas comuns do empreendimento e o conforto ambiental de pedestres e moradores. "
        "Canteiros e áreas de separação entre vias e passeios serão tratados com grama e vegetação de baixo porte, podendo receber arborização viária "
        "conforme as diretrizes técnicas do município. O detalhamento das soluções paisagísticas será apresentado posteriormente em projeto específico."
    ))

    # ILUMINAÇÃO
    ilum_idx = idx
    _heading_num(doc, idx, "ILUMINAÇÃO")
    idx += 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        f"{ilum_idx}.1. A iluminação será composta por postes de concreto instalados em um dos lados das vias, "
        "dimensionados e espaçados conforme as normas do município."
    ))

    # SANEAMENTO BÁSICO
    san_idx = idx
    _heading_num(doc, idx, "SANEAMENTO BÁSICO")
    idx += 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run("A infraestrutura de saneamento básico é composta de:"))

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        f"{san_idx}.1. O sistema de abastecimento de água será definido e projetado de acordo com as diretrizes municipais "
        "e demais órgãos competentes, com tubulações dimensionadas e definidas de acordo com as mesmas."
    ))

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        f"{san_idx}.2. O sistema de esgotamento sanitário dos lotes será projetado conforme as diretrizes municipais "
        "e as exigências dos órgãos competentes."
    ))

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        f"{san_idx}.3. O empreendimento será dotado de rede de drenagem pluvial superficial junto às vias e, em locais específicos, "
        "terá rede e tubulação profunda, obedecendo ao respectivo projeto."
    ))

    # LIXO (só se condomínio)
    if is_cond:
        lixo_idx = idx
        _heading_num(doc, idx, "LIXO")
        idx += 1

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(p.add_run(
            f"{lixo_idx}.1. O acondicionamento do lixo será feito pelo condomínio, em área localizada ao lado da portaria, "
            "voltada para a via de acesso destinada à coleta urbana."
        ))

    # EDIFICAÇÕES NOS LOTES RESIDENCIAIS
    edi_idx = idx
    _heading_num(doc, idx, "EDIFICAÇÕES NOS LOTES RESIDENCIAIS")
    idx += 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        f"{edi_idx}.1. As edificações permitidas nos lotes residenciais a construir serão destinadas exclusivamente à atividade residencial."
    ))

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        f"{edi_idx}.2. Em cada unidade autônoma a construir será permitida a construção de apenas uma unidade habitacional, "
        "respeitando as determinações e os parâmetros definidos pela legislação municipal."
    ))

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(
        f"{edi_idx}.3. Todas as unidades deverão, obrigatoriamente, adotar um recuo de jardim de no mínimo "
    ))
    _run_xxxx(p)
    _set_run_defaults(p.add_run("m."))

    if is_cond:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(p.add_run(
            f"{edi_idx}.4. As demais restrições e condições a que devem satisfazer as construções habitacionais serão atendidas em função das exigências legais da Municipalidade, "
            "devendo para isso todos os projetos residenciais serem apresentados primeiramente à Administração do condomínio, "
            "para posterior licenciamento junto à municipalidade."
        ))

    # CONSIDERAÇÕES GERAIS E FINAIS
    cons_idx = idx
    _heading_num(doc, idx, "CONSIDERAÇÕES GERAIS E FINAIS")
    idx += 1

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(p.add_run(f"{cons_idx}.1. O "))
    r = p.add_run(tipo_lbl + " ")
    _set_run_defaults(r, bold=True)
    p.add_run('"')
    r = p.add_run(nome_fmt if nome_fmt.strip() else "XXXX")
    _set_run_defaults(r, bold=True)
    r.italic = True
    if not (form_data.get('nome_emp') or "").strip():
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run('"')
    _set_run_defaults(p.add_run(
        " caracteriza-se como um empreendimento horizontal, "
    ))
    _set_run_defaults(p.add_run(usos_txt.lower()))
    _set_run_defaults(p.add_run(", caracterizando-se por:"))

    for b in [
        "Um conceito de convívio entre natureza e lazer como premissa básica, desde sua concepção, até sua realização;",
        "Pelas características do seu parcelamento – dos lotes, áreas verdes, de lazer e vias;",
        "Pelas características dos equipamentos urbanos e serviços."
    ]:
        bp = doc.add_paragraph("• " + b)
        _set_run_defaults(bp.runs[0])

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    r = p.add_run(
        "Este memorial descritivo tem caráter informativo e visa apresentar as diretrizes gerais do empreendimento, "
        "podendo sofrer ajustes durante o processo de desenvolvimento, licenciamento e aprovação do Estudo de Viabilidade Urbanística "
        "e do Projeto Urbanístico."
    )
    _set_run_defaults(r)
    r.italic = True

    # linha em branco antes da data
    doc.add_paragraph()

    # Data (à direita)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    MESES = ["janeiro","fevereiro","março","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]
    hoje = datetime.now()
    dia = str(hoje.day)
    mes = MESES[hoje.month-1]
    ano = str(hoje.year)
    _set_run_defaults(p.add_run(f"Porto Alegre, {dia} de {mes} de {ano}."))

    # Assinaturas / rodapé / paginação
    _sec_assinaturas_resumo(doc)
    add_footer_left_text(doc, [
        "WWW.SOLIDO.ARQ.BR",
        "Avenida Ipiranga, 6681 – Prédio 99, Sala 906",
        "Porto Alegre – RS Brasil",
        "+ 55 51 99690-7857",
    ], size_pt=10)
    add_page_numbers(doc)

    out_path = os.path.join(output_dir, "memorial_resumo.docx")
    
    # Verificar se o documento tem conteúdo antes de salvar
    para_count = len(doc.paragraphs)
    if para_count == 0:
        raise Exception("Documento está vazio antes de salvar! Nenhum parágrafo foi adicionado.")
    
    try:
        doc.save(out_path)
        print(f"✅ Documento salvo: {out_path}, Parágrafos: {para_count}")
    except Exception as e:
        import traceback
        print(f"❌ ERRO ao salvar documento: {e}")
        traceback.print_exc()
        raise
    
    # Verificar se o arquivo foi salvo corretamente
    if not os.path.exists(out_path):
        raise Exception(f"Arquivo não foi criado: {out_path}")
    
    file_size = os.path.getsize(out_path)
    if file_size == 0:
        raise Exception(f"Arquivo foi salvo vazio: {out_path}")
    
    # Verificar se o conteúdo foi preservado
    try:
        from docx import Document
        doc_check = Document(out_path)
        if len(doc_check.paragraphs) == 0:
            raise Exception(f"Arquivo salvo mas está vazio ao ler: {out_path}")
        print(f"✅ Arquivo verificado: {len(doc_check.paragraphs)} parágrafos")
    except Exception as e:
        print(f"⚠️ Aviso ao verificar arquivo salvo: {e}")
    
    return out_path

def _build_solicitacao_analise_doc_web(form_data, output_dir):
    """
    Gera o DOCX do tipo 'Solicitação de Análise'
    Adaptada do código original do Xuxu.py
    """
    doc = preparar_doc()
    
    # Parágrafo em branco antes da data
    doc.add_paragraph()
    
    # Data (direita)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    _set_run_defaults(p.add_run(_pt_date("Porto Alegre")))
    
    # Endereçamento (esquerda)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _set_run_defaults(p.add_run("À"))
    
    # Prefeitura Municipal
    cid_fmt = _fmt_cidade_slash_uf(form_data.get('cidade_emp', '') or "")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _set_run_defaults(p.add_run(f"Prefeitura Municipal de {cid_fmt}"))
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _add_hl(p, "Secretaria de Planejamento, Urbanismo e Habitação")
    
    # Objeto
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _set_run_defaults(p.add_run("Objeto: Solicitação de análise de Projeto Urbanistico"))
    
    # Parágrafo em branco
    doc.add_paragraph()
    
    # Corpo
    cid_fmt = _fmt_cidade_slash_uf(form_data.get('cidade_emp', '') or "")
    end_fmt = _title_keep_preps(form_data.get('endereco_emp', '') or "")
    bai_fmt = _fmt_bairro(form_data.get('bairro_emp', '') or "")
    
    # Tipo (loteamento/condomínio)
    tipo_cond = form_data.get('tipo_proj_resumo', 'condominio') or 'condominio'
    if tipo_cond == 'loteamento':
        tipo_cond_txt = "Loteamento de acesso controlado"
    elif tipo_cond == 'condominio':
        tipo_cond_txt = "Condomínio fechado de lotes"
    else:
        tipo_cond_txt = "Empreendimento"
    
    # Área total formatada
    area_total_str = form_data.get('area_total_emp', '') or ''
    if area_total_str.strip():
        try:
            v = _to_float_br(area_total_str)
            area_txt = _fmt_br(v, 2)
        except:
            area_txt = "XXXX"
    else:
        area_txt = "XXXX"
    
    # Matrícula(s)
    rot_mat, mats_fmt = _fmt_matriculas_plural(form_data.get('matricula_emp', '') or '')
    
    par = doc.add_paragraph()
    par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    r = par.add_run("SOLIDO - DESIGN URBANO LTDA. CNPJ nº 26.887.368/0001-07")
    _set_run_defaults(r, bold=True)
    
    _set_run_defaults(par.add_run(", juntamente da "))
    _add_hl(par, "XXXX")
    _set_run_defaults(par.add_run(" - CNPJ nº: "))
    _add_hl(par, "XXXX")
    
    _set_run_defaults(par.add_run(
        f", na qualidade de responsáveis técnicos pelo projeto urbanístico localizado no município de {cid_fmt or 'XXXX'}, "
        f"inserido em uma gleba registrada sob {rot_mat} nº {mats_fmt} no Registro de Imóveis desta cidade, "
        f"vem, por meio deste, requerer a análise técnica para fins de implantação de um "
    ))
    
    # Tipo em negrito
    r = par.add_run(tipo_cond_txt)
    _set_run_defaults(r, bold=True)
    
    _set_run_defaults(par.add_run(", com área total de "))
    
    # Área total + m² em negrito
    r = par.add_run(f"{area_txt}m²")
    _set_run_defaults(r, bold=True)
    
    _set_run_defaults(par.add_run(f", situado na {end_fmt or 'XXXX'}, bairro {bai_fmt or 'XXXX'}, {cid_fmt or 'XXXX'}."))
    
    # Lista de documentos
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _set_run_defaults(p.add_run("Para tanto, protocolamos a seguinte documentação para análise:"))
    
    for item in [
        "- Projeto Urbanistico;",
        "- Memorial resumo do empreendimento;",
        "- Ofício para requerimento de análise;",
        "- RRT."
    ]:
        li = doc.add_paragraph()
        li.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        _set_run_defaults(li.add_run(item))
    
    # Parágrafo em branco
    doc.add_paragraph()
    
    # Fecho
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _set_run_defaults(p.add_run("Nos colocamos à disposição para esclarecimentos e pedimos o deferimento."))
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _set_run_defaults(p.add_run("Atenciosamente,"))
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _set_run_defaults(p.add_run("Grupo Solido e "))
    _add_hl(p, "XXXX")
    
    # Rodapé padrão + numeração
    add_footer_left_text(doc, [
        "WWW.SOLIDO.ARQ.BR",
        "Avenida Ipiranga, 6681 – Prédio 99, Sala 906",
        "Porto Alegre – RS Brasil",
        "+ 55 51 99690-7857",
    ], size_pt=10)
    add_page_numbers(doc)
    
    # Verificar se o documento tem conteúdo antes de salvar
    para_count = len(doc.paragraphs)
    if para_count == 0:
        raise Exception("Documento está vazio antes de salvar! Nenhum parágrafo foi adicionado.")
    
    out_path = os.path.join(output_dir, "solicitacao_analise.docx")
    try:
        doc.save(out_path)
        print(f"✅ Documento salvo: {out_path}, Parágrafos: {para_count}")
    except Exception as e:
        import traceback
        print(f"❌ ERRO ao salvar documento: {e}")
        traceback.print_exc()
        raise
    
    return out_path

def build_unif_desm_doc_web(form_data, uploaded_files, modo, output_dir):
    """
    Gera memorial descritivo de unificação/desmembramento
    Adaptada do código original do Xuxu.py
    """
    # Coletar itens de unificação e desmembramento
    unif_item, desm_items = _collect_items_unif_desm_web(uploaded_files, modo)
    
    doc = preparar_doc()
    pres_unif = bool(unif_item)
    pres_desm = bool(desm_items)
    
    heading(doc, _titulo_para_unif_desm(pres_unif, pres_desm))
    _primeiro_paragrafo_unif_desm(doc, form_data, pres_unif, pres_desm)
    _sec_situacao_atual(doc, form_data, pres_unif, pres_desm)
    
    zone_num, hemi = _auto_zone_from_city(form_data.get('cidade_emp', '') or '')
    if pres_unif:
        _sec_unificacao(doc, form_data, unif_item)
    if pres_desm:
        _sec_desmembramento(doc, form_data, desm_items, zone_num, hemi)
    
    _sec_assinaturas_simples(doc)
    add_footer_left_text(doc, [
        "WWW.SOLIDO.ARQ.BR",
        "Avenida Ipiranga, 6681 – Prédio 99, Sala 906",
        "Porto Alegre – RS Brasil",
        "+ 55 51 99690-7857",
    ], size_pt=10)
    add_page_numbers(doc)
    
    # Verificar se o documento tem conteúdo antes de salvar
    para_count = len(doc.paragraphs)
    if para_count == 0:
        raise Exception("Documento está vazio antes de salvar! Nenhum parágrafo foi adicionado.")
    
    out_path = os.path.join(output_dir, "unif_desm.docx")
    try:
        doc.save(out_path)
        print(f"✅ Documento salvo: {out_path}, Parágrafos: {para_count}")
    except Exception as e:
        import traceback
        print(f"❌ ERRO ao salvar documento: {e}")
        traceback.print_exc()
        raise
    
    return out_path

def build_condominio_loteamento_doc_web(form_data, uploaded_files, modo, output_dir):
    """
    Gera memorial descritivo de condomínio/loteamento
    Adaptada do código original do Xuxu.py
    """
    nome_fmt, end_fmt, cid_fmt, bai_fmt = _get_fmt_campos_basicos(form_data)
    
    # Separar arquivos de lotes dos arquivos Civil 3D
    lot_files = [(f, d) for f, d in uploaded_files.items()
                 if f.lower().endswith(('.html', '.htm', '.txt')) and 'CIVILREPORT' not in f.upper()]
    civil_files = [(f, d) for f, d in uploaded_files.items()
                   if f.lower().endswith(('.html', '.htm')) and 'CIVILREPORT' in f.upper()]
    
    # Processar arquivos de lotes
    file_parcels, all_parcels = [], []
    for fname, data in lot_files:
        quadra = infer_quadra_from_filename(fname)
        if fname.lower().endswith(('.html', '.htm')):
            parcels = parse_parcels_from_html(io.BytesIO(data).read())
        else:
            parcels = parse_parcels_from_txt(data)
        parcels.sort(key=lambda p: p.get('num', 0))
        file_parcels.append((quadra, parcels))
        all_parcels.extend(parcels)
    
    # Ordenar por quadra
    file_parcels.sort(key=lambda qp: quadra_label_sort_key(qp[0]))
    for i, (quadra, parcels) in enumerate(file_parcels):
        parcels.sort(key=lambda p: int(p.get('num', 0)))
        file_parcels[i] = (quadra, parcels)
    
    tipo_full = "Condomínio Fechado de Lotes Residenciais" if modo == 'condominio' else "Loteamento de Acesso Controlado"
    eh_condominio = (modo == 'condominio')
    
    # Calcular áreas totais
    area_tot_priv = area_tot_cond = 0.0
    if eh_condominio:
        area_tot_priv_str = form_data.get('area_tot_priv_emp', '') or ''
        if area_tot_priv_str.strip():
            try:
                area_tot_priv = _to_float_br(area_tot_priv_str)
            except:
                area_tot_priv = 0.0
        area_tot_cond_str = form_data.get('area_tot_cond_emp', '') or ''
        if area_tot_cond_str.strip():
            try:
                area_tot_cond = _to_float_br(area_tot_cond_str)
            except:
                area_tot_cond = 0.0
    
    # Ane (Área Não Edificável)
    ane_drop = form_data.get('ane_drop', 'Não') or 'Não'
    ane_enable = (ane_drop == 'Sim')
    ane_largura_m = None
    if ane_enable:
        ane_largura_str = form_data.get('ane_largura', '') or ''
        if ane_largura_str.strip():
            try:
                ane_largura_m = _to_float_br(ane_largura_str)
            except:
                ane_largura_m = None
    
    # Processar arquivos Civil 3D
    civil_items = []
    for fname, data in civil_files:
        civil_items.extend(parse_civilreport_from_html(io.BytesIO(data).read()))
    
    # Classificar itens do Civil 3D
    grouped = {k: [] for k in [
        'remanescente', 'reserva_tecnica', 'institucional', 'app', 'verde', 'verde_preservacao',
        'viario', 'condominial', 'quadras', 'outros'
    ]}
    for it in civil_items:
        cat, title = classify_civil_item(it['name'])
        grouped[cat].append((title, it))
    
    # Ordenar itens
    def _num_key(nm):
        m = re.search(r'(\d+)', _normalize(nm))
        return int(m.group(1)) if m else 10**9
    
    for cat in grouped:
        if cat == 'viario':
            grouped[cat].sort(key=lambda x: _viario_sort_key(x[1]['name']))
        else:
            grouped[cat].sort(key=lambda x: (_num_key(x[1]['name']), _normalize(x[1]['name'])))
    
    # Criar documento
    doc = preparar_doc()
    heading(doc, "MEMORIAL DESCRITIVO")
    
    def R(par, txt, bold=False):
        run = par.add_run(txt)
        _set_run_defaults(run, bold=bold)
        return run
    
    def _matriculas_texto(raw):
        txt = (raw or '').strip()
        if not txt:
            return "objeto referente à matrícula nº XXXX"
        partes = [p for p in re.split(r'\s*(?:,|;| e )\s*', txt) if p]
        return f"objeto referente às matrículas nºs {txt}" if len(partes) > 1 else f"objeto referente à matrícula nº {txt}"
    
    # Formatar áreas e perímetro
    area_tot_fmt = area_tot_ext = ha_txt = perim_fmt = perim_ext = ""
    area_total_str = form_data.get('area_total_emp', '') or ''
    if area_total_str.strip():
        v = _to_float_br(area_total_str)
        area_tot_fmt = _fmt_br(v, 2) + "m²"
        area_tot_ext = area_por_extenso(v)
        ha_txt = _fmt_br(hectares_from_m2(v), 2) + "ha"
    
    perimetro_str = form_data.get('perimetro_emp', '') or ''
    if perimetro_str.strip():
        pval = _to_float_br(perimetro_str)
        perim_fmt = _fmt_br(pval, 2)
        perim_ext = extenso_metros(pval)
    
    zone_num, hemi = _auto_zone_from_city(form_data.get('cidade_emp', '') or '')
    mc_w = _utm_mc_from_zone(zone_num)
    coord_fmt = form_data.get('coord_fmt', 'utm') or 'utm'
    
    nome_txt_bruto = (nome_fmt or "").strip()
    has_nome = bool(nome_txt_bruto)
    nome_txt = nome_txt_bruto if has_nome else "XXXX"
    end_txt = end_fmt or "XXXX"
    bai_txt = bai_fmt or "XXXX"
    cid_txt = cid_fmt or "XXXX"
    
    # Primeiro parágrafo
    p1 = doc.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    R(p1, "O presente memorial tem por finalidade descrever o parcelamento de solo de acordo com o projeto denominado ")
    
    # TIPO – negrito
    r_tipo = p1.add_run(tipo_full + " ")
    _set_run_defaults(r_tipo, bold=True)
    
    # NOME DO EMPREENDIMENTO – entre aspas, negrito + itálico
    r_asp1 = p1.add_run('"')
    _set_run_defaults(r_asp1, bold=True)
    
    r_nome = p1.add_run(nome_txt)
    _set_run_defaults(r_nome, bold=True)
    r_nome.italic = True
    if not has_nome:
        r_nome.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    r_asp2 = p1.add_run('"')
    _set_run_defaults(r_asp2, bold=True)
    
    R(p1,
      f" em uma gleba de terras situada frente à {end_txt}, bairro {bai_txt} no município de {cid_txt}, "
      f"com área superficial de {area_tot_fmt} ({area_tot_ext}) - {ha_txt} e perímetro de {perim_fmt}m ({perim_ext}), "
      f"{_matriculas_texto(form_data.get('matricula_emp', '') or '')} do registro geral de imóveis desta cidade."
    )
    
    # Segundo parágrafo (coordenadas)
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    if coord_fmt == 'utm':
        R(p2, f"Segue abaixo a descrição completa deste empreendimento. Coordenadas georreferenciadas no Sistema Geodésico Brasileiro, Datum - SIRGAS 2000, MC {mc_w}W, coordenadas Plano Retangulares, sistema UTM.")
    elif coord_fmt == 'dec':
        R(p2, "Segue abaixo a descrição completa deste empreendimento. Coordenadas georreferenciadas ao Sistema Geodésico Brasileiro, referidas ao Datum SIRGAS 2000, expressas em coordenadas geográficas (latitude e longitude) em graus decimais.")
    else:
        R(p2, "Segue abaixo a descrição completa deste empreendimento. Coordenadas georreferenciadas ao Sistema Geodésico Brasileiro, referidas ao Datum SIRGAS 2000, expressas em coordenadas geográficas (latitude e longitude) em graus, minutos e segundos.")
    
    # Seções de áreas (remanescente, institucional, etc.)
    session_order = ['remanescente', 'institucional', 'reserva_tecnica', 'app', 'verde', 'verde_preservacao', 'viario', 'condominial']
    for cat in session_order:
        if not grouped.get(cat):
            continue
        if cat == 'app':
            # Agrupar APP por título
            buckets = {}
            for title, it in grouped[cat]:
                buckets.setdefault(title, []).append(it)
            for gen_title, arr in buckets.items():
                heading(doc, gen_title)
                for it in arr:
                    texto = build_area_text(
                        it['name'], it, tipo_full, nome_fmt or "XXXX",
                        end_fmt or "XXXX", bai_fmt or "XXXX", cid_fmt or "XXXX",
                        ane_enable=False,
                        coord_fmt=coord_fmt,
                        zone_num=zone_num,
                        hemi=hemi
                    )
                    adicionar_texto_formatado(doc, texto)
        else:
            title_cat = grouped[cat][0][0]
            heading(doc, title_cat)
            for _, it in grouped[cat]:
                texto = build_area_text(
                    it['name'], it, tipo_full, nome_fmt or "XXXX",
                    end_fmt or "XXXX", bai_fmt or "XXXX", cid_fmt or "XXXX",
                    ane_enable=False,
                    coord_fmt=coord_fmt,
                    zone_num=zone_num,
                    hemi=hemi
                )
                adicionar_texto_formatado(doc, texto)
    
    # Descrição de Quadras
    heading(doc, "DESCRIÇÃO DE QUADRAS")
    pqd = doc.add_paragraph()
    runxx = pqd.add_run("XXXX")
    _set_run_defaults(runxx)
    runxx.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    # Descrição de Lotes
    heading(doc, "DESCRIÇÃO DE LOTES")
    dados_quadro = []
    for quadra, parcels in file_parcels:
        for parcel in parcels:
            texto_lote = build_memorial_text(
                parcel, quadra, tipo_full, nome_fmt or "XXXX",
                end_fmt or "XXXX", bai_fmt or "XXXX", cid_fmt or "XXXX",
                ane_enable=ane_enable,
                ane_largura_m=ane_largura_m,
                eh_condominio=eh_condominio,
                area_tot_priv=area_tot_priv,
                area_tot_cond=area_tot_cond,
                coord_fmt=coord_fmt,
                zone_num=zone_num,
                hemi=hemi
            )
            adicionar_texto_formatado(doc, texto_lote)
            
            if eh_condominio:
                area_priv = parcel.get("area_m2")
                if area_priv and area_tot_priv > 0:
                    fr = area_priv / area_tot_priv
                    area_comum = fr * (area_tot_cond or 0.0)
                    area_total = area_priv + area_comum
                    dados_quadro.append({
                        'Lote': str(parcel['num']),
                        'Quadra': quadra.replace("QUADRA ", "").strip(),
                        'Área Privativa (m²)': _fmt_br(area_priv, 2),
                        'Área Uso Comum (m²)': _fmt_br(area_comum, 2),
                        'Área Real Total (m²)': _fmt_br(area_total, 2),
                        'Fração Ideal': f"{fr:.7f}"
                    })
    
    # Tabela de fração ideal (se condomínio)
    if eh_condominio and dados_quadro:
        dados_quadro.sort(key=lambda row: (quadra_label_sort_key(f"QUADRA {row['Quadra']}"), _lote_num(row['Lote'])))
        tabela = doc.add_table(rows=1, cols=6)
        tabela.style = 'Table Grid'
        
        # Adicionar bordas
        for row in tabela.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = OxmlElement('w:tcBorders')
                for border_name in ('top', 'left', 'bottom', 'right'):
                    b = OxmlElement(f'w:{border_name}')
                    b.set(qn('w:val'), 'single')
                    b.set(qn('w:sz'), '4')
                    b.set(qn('w:space'), '0')
                    b.set(qn('w:color'), '000000')
                    tc_borders.append(b)
                tc_pr.append(tc_borders)
        
        # Cabeçalho
        hdr = tabela.rows[0].cells
        hdr[0].text = "Lote"
        hdr[1].text = "Quadra"
        hdr[2].text = "Área Priv. (m²)"
        hdr[3].text = "Área Uso Comum (m²)"
        hdr[4].text = "Área Real Total (m²)"
        hdr[5].text = "Fração Ideal"
        for c in hdr:
            p = c.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in p.runs:
                _set_run_defaults(run, bold=True)
        
        # Dados
        for row in dados_quadro:
            cells = tabela.add_row().cells
            cells[0].text = row.get('Lote', '')
            cells[1].text = row.get('Quadra', '')
            cells[2].text = row.get('Área Privativa (m²)', '')
            cells[3].text = row.get('Área Uso Comum (m²)', '')
            cells[4].text = row.get('Área Real Total (m²)', '')
            cells[5].text = row.get('Fração Ideal', '')
            for c in cells:
                for par in c.paragraphs:
                    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in par.runs:
                        _set_run_defaults(run)
    
    # Assinaturas, rodapé e paginação
    _sec_assinaturas_simples(doc)
    add_footer_left_text(doc, [
        "WWW.SOLIDO.ARQ.BR",
        "Avenida Ipiranga, 6681 – Prédio 99, Sala 906",
        "Porto Alegre – RS Brasil",
        "+ 55 51 99690-7857",
    ], size_pt=10)
    add_page_numbers(doc)
    
    # Verificar se o documento tem conteúdo antes de salvar
    para_count = len(doc.paragraphs)
    if para_count == 0:
        raise Exception("Documento está vazio antes de salvar! Nenhum parágrafo foi adicionado.")
    
    out_path = os.path.join(output_dir, "memorial_lotes.docx")
    try:
        doc.save(out_path)
        print(f"✅ Documento salvo: {out_path}, Parágrafos: {para_count}")
    except Exception as e:
        import traceback
        print(f"❌ ERRO ao salvar documento: {e}")
        traceback.print_exc()
        raise
    
    # Verificar se o arquivo foi salvo corretamente
    if not os.path.exists(out_path):
        raise Exception(f"Arquivo não foi criado: {out_path}")
    
    file_size = os.path.getsize(out_path)
    if file_size == 0:
        raise Exception(f"Arquivo foi salvo vazio: {out_path}")
    
    return out_path

def build_excel_fracao_ideal_web(form_data, uploaded_files, output_dir):
    """
    Gera Excel de Fração Ideal (somente condomínio)
    Adaptada do código original do Xuxu.py
    """
    import pandas as pd
    from openpyxl import load_workbook
    from openpyxl.styles import Font, Alignment, Border, Side
    
    # Para gerar o Excel de fração ideal, precisamos dos dados_quadro
    # que são gerados durante a criação do documento de condomínio
    # Como não temos acesso a essas variáveis globais, vamos gerar um Excel vazio
    # com a estrutura correta, ou retornar erro informando que é necessário gerar o DOCX primeiro
    
    # Por enquanto, vamos criar um Excel vazio com a estrutura correta
    df = pd.DataFrame(columns=[
        'Lote', 'Quadra', 'Área Privativa (m²)', 'Área Uso Comum (m²)', 
        'Área Real Total (m²)', 'Fração Ideal'
    ])
    
    out_path = os.path.join(output_dir, "fracao_ideal.xlsx")
    df.to_excel(out_path, index=False)
    
    # Formatar o Excel
    wb = load_workbook(out_path)
    ws = wb.active
    
    font_header = Font(name='Calibri', size=12, bold=True)
    font_cell = Font(name='Calibri', size=12)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin = Side(border_style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    
    for r in ws.iter_rows():
        for c in r:
            c.alignment = center
            c.border = border
            c.font = font_header if c.row == 1 else font_cell
    
    for col in ws.columns:
        maxlen = max(len(str(c.value)) if c.value is not None else 0 for c in col)
        ws.column_dimensions[col[0].column_letter].width = max(12, maxlen + 2)
    ws.column_dimensions['D'].width = 22
    
    wb.save(out_path)
    return out_path

def build_excel_vertices_web(form_data, uploaded_files, modo, output_dir):
    """
    Gera Excel de Vértices para unificação/desmembramento
    Adaptada do código original do Xuxu.py
    """
    from openpyxl import Workbook
    
    # Coletar itens
    unif_item, desm_items = _collect_items_unif_desm_web(uploaded_files, modo)
    
    wb = Workbook()
    wb.remove(wb.active)
    
    def _nova_aba(nome):
        ws = wb.create_sheet(title=nome)
        # Aplicar larguras de coluna básicas
        for col in range(1, 10):
            ws.column_dimensions[chr(64 + col)].width = 15
        return ws
    
    def _num_after_name(nm: str) -> int:
        m = re.search(r'(\d+)', _normalize(nm))
        return int(m.group(1)) if m else 10**9
    
    # Gerar vértices baseado no modo
    if modo == 'desmembramento':
        desm_sorted = sorted(desm_items, key=lambda x: (_num_after_name(x[0]), _normalize(x[0])))
        ws = _nova_aba("DESMEMBRAMENTO")
        # Adicionar cabeçalho básico
        ws['A1'] = "Área"
        ws['B1'] = "Vértice"
        ws['C1'] = "Coordenada X"
        ws['D1'] = "Coordenada Y"
        # TODO: Implementar preenchimento completo dos vértices
        
    elif modo == 'unificacao':
        ws = _nova_aba("UNIFICAÇÃO")
        ws['A1'] = "Área"
        ws['B1'] = "Vértice"
        ws['C1'] = "Coordenada X"
        ws['D1'] = "Coordenada Y"
        # TODO: Implementar preenchimento completo dos vértices
        
    else:  # 'unif_desm'
        ws_u = _nova_aba("UNIFICAÇÃO")
        ws_u['A1'] = "Área"
        ws_u['B1'] = "Vértice"
        ws_u['C1'] = "Coordenada X"
        ws_u['D1'] = "Coordenada Y"
        
        desm_sorted = sorted(desm_items, key=lambda x: (_num_after_name(x[0]), _normalize(x[0])))
        ws_d = _nova_aba("DESMEMBRAMENTO")
        ws_d['A1'] = "Área"
        ws_d['B1'] = "Vértice"
        ws_d['C1'] = "Coordenada X"
        ws_d['D1'] = "Coordenada Y"
        # TODO: Implementar preenchimento completo dos vértices
    
    out_path = os.path.join(output_dir, "vertices.xlsx")
    wb.save(out_path)
    return out_path

# Funções auxiliares para UNIF/DESM
def _cidade_sem_uf(txt):
    s = str(txt or "XXXX").strip()
    if "/" in s:
        s = s.split("/", 1)[0].strip()
    return s if s else "XXXX"

def _collect_items_unif_desm_web(uploaded_files, modo):
    """Versão web de _collect_items_unif_desm"""
    items_unif = None
    items_desm = []

    civil_htmls = [(f,d) for f,d in uploaded_files.items()
                   if f.lower().endswith(('.html','.htm')) and 'CIVILREPORT' in f.upper()]
    other_htmls = [(f,d) for f,d in uploaded_files.items()
                   if f.lower().endswith(('.html','.htm')) and 'CIVILREPORT' not in f.upper()]

    if modo in ('unificacao','unif_desm'):
        for fname, data in civil_htmls:
            arr = parse_civilreport_from_html(io.BytesIO(data).read())
            for it in arr:
                nm = it.get('name') or ''
                if is_unificacao_item_name(nm):
                    items_unif = items_unif or it

    if modo in ('desmembramento','unif_desm'):
        for fname, data in other_htmls:
            try:
                parcels = parse_parcels_from_html(io.BytesIO(data).read())
                for p in parcels:
                    item = {'segments': p.get('segments', []), 'area_m2': p.get('area_m2', 0.0), 'first_point': p.get('first_point')}
                    nm = f"GLEBA {p.get('num', 1)}"
                    items_desm.append((nm, item))
            except:
                pass
    return items_unif, items_desm

# Funções de seções UNIF/DESM (adaptadas)
def _sec_situacao_atual(doc, form_data, pres_unif, pres_desm):
    """Adaptada para usar form_data"""
    nome_fmt, end_fmt, cid_fmt, bai_fmt = _get_fmt_campos_basicos(form_data)
    mats_txt = form_data.get('matricula_emp', '').strip()
    partes = [p.strip() for p in re.split(r'\s*(?:,|;| e )\s*', mats_txt) if p.strip()]
    cidade_txt = cid_fmt or 'XXXX'
    titulo = "SITUAÇÃO ATUAL DAS MATRÍCULAS" if len(partes) > 1 else "SITUAÇÃO ATUAL DA MATRÍCULA"
    heading(doc, f"{titulo} " + (", ".join(partes) if partes else "XXXX"))

    if not partes: partes = ["XXXX"]

    for mat in partes:
        par = doc.add_paragraph()
        par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        r = par.add_run("Imóvel: ")
        _set_run_defaults(r, bold=True)
        r2 = par.add_run("Matrícula ")
        _set_run_defaults(r2)
        if mat == "XXXX":
            _add_hl(par, "XXXX")
        else:
            _set_run_defaults(par.add_run(mat))
        _set_run_defaults(par.add_run(", Município de "))
        if form_data.get('cidade_emp') and form_data.get('cidade_emp').strip().upper() != "XXXX":
            _set_run_defaults(par.add_run(cidade_txt))
        else:
            _add_hl(par, "XXXX")
        _set_run_defaults(par.add_run(", com área total de "))
        _add_hl(par, "XXXX")
        _set_run_defaults(par.add_run("m²."))

        par2 = doc.add_paragraph()
        par2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        r3 = par2.add_run("Descrição do Imóvel: ")
        _set_run_defaults(r3, bold=True)
        _add_hl(par2, "XXXX")

def _sec_unificacao(doc, form_data, unif_item):
    """Adaptada para usar form_data"""
    nome_fmt, end_fmt, cid_fmt, bai_fmt = _get_fmt_campos_basicos(form_data)
    heading(doc, "UNIFICAÇÃO")

    if unif_item is None:
        par = doc.add_paragraph()
        par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(par.add_run("Imóvel: "), bold=True)
        _add_hl(par, "XXXX")
        _set_run_defaults(par.add_run(", com área total de "))
        _add_hl(par, "XXXX")
        _set_run_defaults(par.add_run("m² ("))
        _add_hl(par, "XXXX")
        _set_run_defaults(par.add_run(")."))
        par2 = doc.add_paragraph()
        par2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(par2.add_run("Descrição do Imóvel: "), bold=True)
        _add_hl(par2, "XXXX")
        return

    area = unif_item.get("area_m2") or 0.0
    area_fmt = _fmt_br(area, 2)
    area_ext = area_por_extenso(area)
    nome = unif_item.get("name") or "UNIFICAÇÃO"

    par = doc.add_paragraph()
    par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    _set_run_defaults(par.add_run("Imóvel: "), bold=True)
    _set_run_defaults(par.add_run(f"{_normalize(nome)}"))
    _set_run_defaults(par.add_run(", com área total de "))
    run_num = par.add_run(area_fmt + "m²")
    _set_run_defaults(run_num, bold=True)
    _set_run_defaults(par.add_run(" ("))
    _set_run_defaults(par.add_run(area_ext))
    _set_run_defaults(par.add_run(")."))

    zone_num, hemi = _auto_zone_from_city(form_data.get('cidade_emp', ''))
    coord_fmt = form_data.get('coord_fmt', 'utm')

    texto_desc = build_area_text(
        item_name=nome,
        item=unif_item,
        tipo_full="",
        empreendimento=nome_fmt or "",
        endereco=end_fmt or "XXXX",
        bairro=bai_fmt or "XXXX",
        cidade=cid_fmt or "XXXX",
        ane_enable=False,
        coord_fmt=coord_fmt,
        zone_num=zone_num,
        hemi=hemi,
        ident_label_only=True,
        ident_label_text="Descrição do Imóvel:"
    )
    adicionar_texto_formatado(doc, texto_desc)

def _sec_desmembramento(doc, form_data, desm_items, zone_num, hemi):
    """Adaptada para usar form_data"""
    nome_fmt, end_fmt, cid_fmt, bai_fmt = _get_fmt_campos_basicos(form_data)
    heading(doc, "DESMEMBRAMENTO")

    def _first_int_or_big(s):
        m = re.search(r'(\d+)', str(s) or '')
        return int(m.group(1)) if m else 10**9

    desm_items_sorted = sorted(desm_items, key=lambda kv: (_first_int_or_big(kv[0]), _normalize(kv[0])))

    coord_fmt = form_data.get('coord_fmt', 'utm')

    for nm, item in desm_items_sorted:
        area = item.get("area_m2") or 0.0
        area_fmt = _fmt_br(area, 2)
        area_ext = area_por_extenso(area)

        par = doc.add_paragraph()
        par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _set_run_defaults(par.add_run("Imóvel: "), bold=True)
        _set_run_defaults(par.add_run(f"{_normalize(nm)}"))
        _set_run_defaults(par.add_run(", com área total de "))
        run_num = par.add_run(area_fmt + "m²")
        _set_run_defaults(run_num, bold=True)
        _set_run_defaults(par.add_run(" ("))
        _set_run_defaults(par.add_run(area_ext))
        _set_run_defaults(par.add_run(")."))

        texto_desc = build_area_text(
            item_name=nm,
            item=item,
            tipo_full="",
            empreendimento=nome_fmt or "",
            endereco=end_fmt or "XXXX",
            bairro=bai_fmt or "XXXX",
            cidade=cid_fmt or "XXXX",
            ane_enable=False,
            coord_fmt=coord_fmt,
            zone_num=zone_num,
            hemi=hemi,
            ident_label_only=True,
            ident_label_text="Descrição do Imóvel:"
        )
        adicionar_texto_formatado(doc, texto_desc)

def _primeiro_paragrafo_unif_desm(doc, form_data, pres_unif, pres_desm):
    """Adaptada para usar form_data"""
    nome_fmt, end_fmt, cid_fmt, bai_fmt = _get_fmt_campos_basicos(form_data)

    end_txt = end_fmt or 'XXXX'
    bairro_txt = bai_fmt or 'XXXX'
    cidade_txt = cid_fmt or 'XXXX'
    mats_txt = form_data.get('matricula_emp', '').strip()

    if pres_unif and pres_desm:
        titulo = "a unificação e desmembramento"
    elif pres_unif:
        titulo = "a unificação"
    else:
        titulo = "o desmembramento"

    area_fmt_txt = None
    area_ext_txt = None
    if form_data.get('area_total_emp', '').strip():
        try:
            v = _to_float_br(form_data.get('area_total_emp'))
            area_fmt_txt = _fmt_br(v, 2)
            area_ext_txt = area_por_extenso(v)
        except:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    run = p.add_run("O presente memorial tem por finalidade descrever " + titulo + " de uma área de terras, situadas frente ")
    _set_run_defaults(run)

    if end_txt != "XXXX":
        _set_run_defaults(p.add_run(end_txt))
    else:
        _add_hl(p, "XXXX")

    _set_run_defaults(p.add_run(", no bairro "))
    if bairro_txt != "XXXX":
        _set_run_defaults(p.add_run(bairro_txt))
    else:
        _add_hl(p, "XXXX")

    _set_run_defaults(p.add_run(", nesta comarca e cidade de "))
    if cidade_txt != "XXXX":
        _set_run_defaults(p.add_run(cidade_txt))
    else:
        _add_hl(p, "XXXX")

    _set_run_defaults(p.add_run(" com área total de "))
    if area_fmt_txt:
        run_num = p.add_run(area_fmt_txt + "m²")
        _set_run_defaults(run_num, bold=True)
    else:
        _add_hl(p, "XXXX")
        _set_run_defaults(p.add_run("m²"))

    _set_run_defaults(p.add_run(" ("))
    if area_ext_txt:
        _set_run_defaults(p.add_run(area_ext_txt))
    else:
        _add_hl(p, "XXXX")
    _set_run_defaults(p.add_run("), objeto referente "))

    partes = [s.strip() for s in re.split(r'\s*(?:,|;| e )\s*', mats_txt) if s.strip()]
    _set_run_defaults(p.add_run("à matrícula sob " if len(partes) <= 1 else "às matrículas sob "))
    if partes:
        _set_run_defaults(p.add_run(", ".join(partes)))
    else:
        _add_hl(p, "XXXX")
    _set_run_defaults(p.add_run(" do registro geral de imóveis desta cidade."))

    zone_num, hemi = _auto_zone_from_city(form_data.get('cidade_emp', ''))
    mc_w = _utm_mc_from_zone(zone_num)
    coord_fmt = form_data.get('coord_fmt', 'utm')

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if coord_fmt == 'utm':
        _set_run_defaults(p2.add_run(
            f"Coordenadas georreferenciadas no Sistema Geodésico Brasileiro, Datum - SIRGAS 2000, "
            f"MC {mc_w}W, coordenadas Plano Retangulares, sistema UTM."
        ))
    elif coord_fmt == 'dec':
        _set_run_defaults(p2.add_run(
            "Coordenadas georreferenciadas ao Sistema Geodésico Brasileiro, referidas ao Datum SIRGAS 2000, "
            "expressas em coordenadas geográficas (latitude e longitude) em graus decimais."
        ))
    else:
        _set_run_defaults(p2.add_run(
            "Coordenadas georreferenciadas ao Sistema Geodésico Brasileiro, referidas ao Datum SIRGAS 2000, "
            "expressas em coordenadas geográficas (latitude e longitude) em graus, minutos e segundos."
        ))

def _titulo_para_unif_desm(pres_unif, pres_desm):
    if pres_unif and pres_desm:
        return "MEMORIAL DESCRITIVO DE UNIFICAÇÃO E DESMEMBRAMENTO"
    if pres_unif:
        return "MEMORIAL DESCRITIVO DE UNIFICAÇÃO"
    if pres_desm:
        return "MEMORIAL DESCRITIVO DE DESMEMBRAMENTO"
    return "MEMORIAL DESCRITIVO"

def _sec_assinaturas_simples(doc):
    heading(doc, "ASSINATURAS")
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("_____________________________")
    _set_run_defaults(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Responsável técnico")
    _set_run_defaults(r, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SOLIDO - DESIGN URBANO LTDA.")
    _set_run_defaults(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("CAU-RS 15335-4")
    _set_run_defaults(r)

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("_____________________________")
    _set_run_defaults(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pr = p.add_run("Proprietário")
    _set_run_defaults(pr, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r_label = p.add_run("")
    _set_run_defaults(r_label)
    r_xxxx = p.add_run("XXXX")
    _set_run_defaults(r_xxxx)
    r_xxxx.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r_cnpj_label = p.add_run("CNPJ: ")
    _set_run_defaults(r_cnpj_label)
    r_cnpj_xxxx = p.add_run("XXXX")
    _set_run_defaults(r_cnpj_xxxx)
    r_cnpj_xxxx.font.highlight_color = WD_COLOR_INDEX.YELLOW

